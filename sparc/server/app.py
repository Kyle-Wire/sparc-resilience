"""
SPARC FastAPI Server
====================

Thin HTTP / WebSocket layer over the existing ``sparc/`` pipeline.
Launched by the Tauri desktop app as a sidecar on ``localhost:8008``,
or manually via ``sparc server --port 8008``.

The server holds loaded project data and model objects in memory so that
the React frontend can query results without disk round-trips, and
pipeline stages stream structured JSON events over a WebSocket instead
of dumping to stdout.
"""

from __future__ import annotations

import asyncio
import os
import shutil
import tempfile
from pathlib import Path
from typing import Any, Literal, Optional

import numpy as np
from pydantic import BaseModel, Field

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Query, Body, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, Response

from sparc.server.state import ServerState
from sparc.server.stream import stream_stage

# ------------------------------------------------------------------
# Application & shared state
# ------------------------------------------------------------------

app = FastAPI(
    title="SPARC Server",
    version="1.0.0",
    docs_url="/docs",
)

# Allow the Tauri webview (and dev server) to reach us
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "tauri://localhost",
        "http://tauri.localhost",
        "https://tauri.localhost",
        "http://localhost:1420",
        "http://localhost:5173",
    ],
    allow_methods=["*"],
    allow_headers=["*"],
)

state = ServerState()

TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "templates"


# ------------------------------------------------------------------
# Physics helpers
# ------------------------------------------------------------------

def _resolve_expected_sign(cfg: dict, variable: str) -> str | None:
    """Return ``'negative'``, ``'positive'``, or ``None`` for *variable*.

    Checks three sources in order:
      1. ``config["physics"]["monotone_constraints"]``  (inline in project.yml)
      2. ``config["physics"]["caps_file"]`` → ``monotonicity.{var}.expected_sign``
      3. Built-in PhysicsPriors defaults (Canopy → negative, etc.)
    """
    physics = cfg.get("physics", {})
    if not isinstance(physics, dict):
        return None

    # 1. monotone_constraints  e.g. {"Pct_Canopy": -1, "NDVI": -1}
    mc = physics.get("monotone_constraints", {})
    if isinstance(mc, dict) and variable in mc:
        val = mc[variable]
        if val == -1 or val == "-1":
            return "negative"
        if val == 1 or val == "1":
            return "positive"
        return None  # 0 = unconstrained

    # 2. caps_file → monotonicity section
    caps_path = physics.get("caps_file")
    if caps_path and os.path.exists(caps_path):
        try:
            import yaml
            with open(caps_path, "r", encoding="utf-8") as fh:
                caps = yaml.safe_load(fh) or {}
            mono = caps.get("monotonicity", {})
            if isinstance(mono, dict) and variable in mono:
                return mono[variable].get("expected_sign") if isinstance(mono[variable], dict) else None
        except Exception:
            pass

    # 3. Built-in PhysicsPriors class (literature defaults)
    try:
        from sparc.interventions.physics_priors import PhysicsPriors
        pp = PhysicsPriors()
        coef = pp.coefficients.get(variable)
        if coef is not None:
            return "negative" if coef.coefficient < 0 else "positive" if coef.coefficient > 0 else None
    except Exception:
        pass

    return None


# ------------------------------------------------------------------
# Scenario GPKG cache (avoids re-reading multi-MB file per slider tick)
# ------------------------------------------------------------------

_scenario_gpkg_cache: dict[str, Any] = {}  # {path_str: GeoDataFrame}


def _load_scenario_gpkg(paths) -> Any:
    """Return the scenario GeoDataFrame (WGS84), cached after first load."""
    for gpkg_name in ("scenario_results.gpkg", "scenario_results_dag.gpkg",
                      "scenario_results_hybrid.gpkg", "scenario_results_reprediction.gpkg"):
        candidate = paths.stage4_dir / gpkg_name
        if candidate.exists():
            key = str(candidate)
            if key in _scenario_gpkg_cache:
                return _scenario_gpkg_cache[key]
            import geopandas as gpd
            gdf = gpd.read_file(candidate)
            _scenario_gpkg_cache[key] = gdf
            return gdf
    return None


# ------------------------------------------------------------------
# Run registry helpers
# ------------------------------------------------------------------

def _attach_registry(config: dict) -> None:
    """Attach a RunRegistry to ``state``; runs migrate_from_disk for legacy runs.

    Failures here are non-fatal — endpoints fall back to disk paths.
    """
    try:
        from sparc.registry import RunRegistry
        from sparc.run.pipeline_paths import PipelinePaths
        paths = PipelinePaths.from_config(config)
        reg = RunRegistry(paths.output_dir, autoload=True)
        # If the manifest is empty (legacy / freshly-loaded run), import what's
        # already on disk so the frontend sees correct availability.
        try:
            reg.migrate_from_disk(paths)
        except Exception as exc:  # noqa: BLE001
            print(f"Warning: registry migration failed: {exc}")
        # Detach previous listener (if any) before swapping registries.
        prev = state.registry
        if prev is not None:
            try:
                prev.remove_register_listener(_on_artifact_registered)
            except Exception:  # noqa: BLE001
                pass
        try:
            reg.add_register_listener(_on_artifact_registered)
        except Exception as exc:  # noqa: BLE001
            print(f"Warning: could not attach artifact listener: {exc}")
        state.registry = reg
    except Exception as exc:  # noqa: BLE001
        print(f"Warning: could not attach RunRegistry: {exc}")
        state.registry = None


# ------------------------------------------------------------------
# Live artifact-event broadcasting
# ------------------------------------------------------------------

# Active /run/stream subscribers — populated when a websocket connects.
# Each entry is an ``asyncio.Queue`` of pending events for that subscriber.
_artifact_subscribers: list[Any] = []


def _on_artifact_registered(entry: Any) -> None:
    """Listener attached to ``RunRegistry.register_artifact``.

    Buffers an ``artifact_written`` event for ``/run/events`` polling and
    fans it out to any live ``/run/stream`` subscribers.
    """
    event = {
        "type": "artifact_written",
        "stage": str(getattr(entry, "stage", "")),
        "artifact_id": getattr(entry, "id", None),
        "kind": getattr(entry, "storage_kind", None),
        "format": getattr(entry, "format", None),
        "size_bytes": getattr(entry, "size_bytes", 0),
        "row_count": getattr(entry, "row_count", None),
        "content_hash": getattr(entry, "sha256", None) or getattr(entry, "blob_sha256", None),
        "written_at": getattr(entry, "written_at", None),
    }
    try:
        state.buffer_event(event)
    except Exception:  # noqa: BLE001
        pass
    # Fan out to live ws subscribers via their queues. Listener may run on a
    # background thread; use ``call_soon_threadsafe`` if a loop is running.
    for q in list(_artifact_subscribers):
        try:
            loop = getattr(q, "_sparc_loop", None)
            if loop is not None and loop.is_running():
                loop.call_soon_threadsafe(q.put_nowait, event)
            else:
                q.put_nowait(event)
        except Exception:  # noqa: BLE001
            pass


def _registry_path(stage: str | int, artifact_id: str) -> Path | None:
    """Look up an artifact's absolute path via the registry. None if missing."""
    reg = state.registry
    if reg is None:
        return None
    entry = reg.lookup(stage, artifact_id)
    if entry is None or entry.partial:
        return None
    p = reg.resolve(entry)
    return p if p.exists() else None


def _open_store():
    """Return an ``ArtifactStore`` bound to the active registry, or raise 400.

    All ``/results/*`` endpoints are db-only: artifacts must live in
    ``artifacts.db``. Disk fallbacks were removed in the v4 refresh.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(400, "No active run registry. Load a project first.")
    from sparc.registry.run_registry import set_active_registry
    from sparc.registry.store import ArtifactStore
    set_active_registry(state.registry)
    return ArtifactStore(state.registry)


def _read_or_404(
    stage: str | int,
    artifact_id: str,
    *,
    hint: str = "",
):
    """DB-only read: ``store.read_any`` or structured 404 if missing.

    The frontend's ``parseMissingArtifact`` consumes the structured
    detail to render an actionable empty-state.
    """
    store = _open_store()
    try:
        if not store.has(stage, artifact_id):
            raise _missing_artifact_response(
                artifact_id=artifact_id, stage=stage, hint=hint,
            )
        return store.read_any(stage, artifact_id)
    finally:
        from sparc.registry.run_registry import set_active_registry
        set_active_registry(None)


def _missing_artifact_response(
    *,
    artifact_id: str,
    stage: str | int,
    expected_paths: list[Path] | None = None,
    hint: str = "",
) -> HTTPException:
    """Return a structured 404 the frontend can render as an actionable empty-state."""
    detail = {
        "error": "missing_artifact",
        "missing_artifact": artifact_id,
        "produced_by_stage": str(stage),
        "expected_path": (str(expected_paths[0]) if expected_paths else None),
        "candidate_paths": [str(p) for p in (expected_paths or [])],
        "hint": hint,
    }
    return HTTPException(status_code=404, detail=detail)


# ------------------------------------------------------------------
# Startup: auto-load project if SPARC_SERVER_PROJECT env var is set
# ------------------------------------------------------------------

@app.on_event("startup")
async def _auto_load_project():
    project_env = os.environ.get("SPARC_SERVER_PROJECT")
    if not project_env:
        return
    resolved = Path(project_env).resolve()
    if not resolved.exists():
        print(f"Warning: SPARC_SERVER_PROJECT file not found: {resolved}")
        return
    try:
        from sparc.config.config import load_config
        import yaml
        with open(resolved, 'r', encoding='utf-8') as fh:
            raw_yaml = yaml.safe_load(fh) or {}
        config = load_config(str(resolved))
        state.project_path = str(resolved)
        state.project_config = config
        state.raw_project_yaml = raw_yaml
        _attach_registry(config)
        _load_data_into_state(config)
        print(f"Auto-loaded project: {resolved}")
    except Exception as exc:
        print(f"Warning: auto-load failed for {resolved}: {exc}")


# ------------------------------------------------------------------
# Health
# ------------------------------------------------------------------

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "project_loaded": state.project_config is not None,
        "project_path": state.project_path,
        "is_running": state.is_running,
        "current_stage": state.current_stage,
        "manifest_loaded": state.registry is not None,
    }


@app.get("/debug/paths")
async def debug_paths():
    """Diagnostic endpoint: show resolved output paths and whether expected files exist."""
    if state.project_config is None:
        return {"error": "No project loaded"}

    from sparc.run.pipeline_paths import PipelinePaths

    try:
        paths = PipelinePaths.from_config(state.project_config)
    except Exception as exc:
        return {"error": f"Cannot resolve paths: {exc}"}

    expected_files = {
        "stage1_gwen_csv": str(paths.stage1_dir / "gwen_variable_importance.csv"),
        "stage1_gwen_json": str(paths.stage1_dir / "gwen_results.json"),
        "stage2_predictions_gpkg": str(paths.stage2_dir / "spatial_cv_predictions.gpkg"),
        "stage3_coefficients": str(paths.stage3_dir / "scenario_coefficients.json"),
        "stage3_dose_response": str(paths.stage3_dir / "dose_response_curves.json"),
        "stage4_scenario_gpkg": str(paths.stage4_dir / "scenario_results.gpkg"),
    }

    return {
        "output_dir": str(paths.output_dir),
        "stage1_dir": str(paths.stage1_dir),
        "stage2_dir": str(paths.stage2_dir),
        "stage3_dir": str(paths.stage3_dir),
        "stage4_dir": str(paths.stage4_dir),
        "final_dir": str(paths.final_dir),
        "files": {k: {"path": v, "exists": Path(v).exists()} for k, v in expected_files.items()},
        "config_output_base_dir": state.project_config.get("output", {}).get("base_dir"),
        "config_paths_output_dir": state.project_config.get("paths", {}).get("output_dir"),
    }


# ------------------------------------------------------------------
# Block export — capture the current view of any chart/map/plot as a
# PNG and register it in the run registry so it shows up in Report.
# ------------------------------------------------------------------

class BlockExportRequest(BaseModel):
    artifact_id: str
    stage: str = "export"
    label: Optional[str] = None
    png_b64: str  # data URL or raw base64

    class Config:
        extra = "ignore"


@app.post("/results/export")
async def export_block(req: BlockExportRequest):
    """Persist a screenshot of a UI block (map / chart / plot) to disk and the registry.

    Body fields:
      - artifact_id: stable id for this export (e.g. "cate_map", "dose_response").
      - stage: registry stage to file under (defaults to "export").
      - label: optional human-readable suffix used in the filename.
      - png_b64: data URL ("data:image/png;base64,...") or raw base64 PNG bytes.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    import base64
    import time
    from sparc.run.pipeline_paths import PipelinePaths

    payload = req.png_b64
    if payload.startswith("data:"):
        try:
            payload = payload.split(",", 1)[1]
        except IndexError:
            raise HTTPException(400, "Malformed data URL in png_b64")
    try:
        raw = base64.b64decode(payload, validate=True)
    except Exception as exc:
        raise HTTPException(400, f"Invalid base64 PNG: {exc}")
    if not raw.startswith(b"\x89PNG\r\n\x1a\n"):
        raise HTTPException(400, "png_b64 does not contain a PNG signature")

    try:
        paths = PipelinePaths.from_config(state.project_config)
    except Exception as exc:
        raise HTTPException(500, f"Cannot resolve paths: {exc}")

    exports_dir = paths.output_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    ts = time.strftime("%Y%m%dT%H%M%S")
    safe_label = "".join(c for c in (req.label or "") if c.isalnum() or c in "-_") or "snapshot"
    fname = f"{req.stage}__{req.artifact_id}__{safe_label}__{ts}.png"
    out_path = exports_dir / fname
    out_path.write_bytes(raw)

    entry_dict: dict[str, Any] = {
        "saved_to": str(out_path),
        "size_bytes": len(raw),
        "registered": False,
    }

    reg = state.registry
    if reg is not None:
        try:
            entry = reg.register_artifact(
                stage=req.stage,
                artifact_id=f"export::{req.artifact_id}::{ts}",
                path=out_path,
                format="png",
                producer="desktop:export",
                consumers=["user:export", "report:figures"],
                metadata={
                    "exported_from": req.artifact_id,
                    "label": req.label or "",
                    "exported_at": ts,
                },
            )
            entry_dict["registered"] = True
            entry_dict["artifact"] = entry.model_dump()
        except Exception as exc:
            entry_dict["registry_error"] = str(exc)

    return entry_dict


# ------------------------------------------------------------------
# Project endpoints
# ------------------------------------------------------------------

@app.post("/project/load")
async def load_project(path: str = Query(..., description="Absolute path to project.yml")):
    """Load a project.yml and return its metadata."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        raise HTTPException(404, f"Project file not found: {resolved}")

    from sparc.config.config import load_config

    # Also keep the raw YAML for the config editor
    import yaml
    with open(resolved, 'r', encoding='utf-8') as fh:
        raw_yaml = yaml.safe_load(fh) or {}

    try:
        config = load_config(str(resolved))
    except Exception as exc:
        raise HTTPException(422, f"Invalid project configuration: {exc}")

    state.project_path = str(resolved)
    state.project_config = config
    state.raw_project_yaml = raw_yaml
    _attach_registry(config)

    # Pre-load data summary if the CSV exists
    data_path = config["data"]["file_path"]
    if os.path.exists(data_path):
        _load_data_into_state(config)

    return {
        "status": "loaded",
        "project": raw_yaml.get("project", {}),
        "columns": list(state.data.columns) if state.data is not None else [],
        "row_count": len(state.data) if state.data is not None else 0,
    }


@app.post("/project/validate")
async def validate_project(path: str = Query(..., description="Absolute path to project.yml")):
    """Validate a project.yml without loading it into state."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        raise HTTPException(404, f"Project file not found: {resolved}")

    from sparc.config.config import load_config

    warnings: list[str] = []
    try:
        config = load_config(str(resolved))
    except Exception as exc:
        return {"valid": False, "error": str(exc), "warnings": []}

    # Check data file
    data_path = config["data"]["file_path"]
    if not os.path.exists(data_path):
        warnings.append(f"Data file not found: {data_path}")
    else:
        import pandas as pd
        df = pd.read_csv(data_path, nrows=5)
        expected = (
            [config["variables"]["target"]]
            + list(config["variables"]["coordinates"])
            + list(config["predictors"]["base_model"])
        )
        missing = [c for c in expected if c not in df.columns]
        if missing:
            warnings.append(f"Missing columns: {missing}")

    # Check physics files
    for key in ("priors_file", "caps_file"):
        fpath = config.get("physics", {}).get(key)
        if fpath and not os.path.exists(fpath):
            warnings.append(f"{key} not found: {fpath}")

    # Check DAG file
    dag_file = config.get("causal", {}).get("dag_file")
    if dag_file and not os.path.exists(dag_file):
        warnings.append(f"dag_file not found: {dag_file}")

    return {"valid": True, "warnings": warnings}


@app.post("/project/init")
async def init_project(
    template: str = Query("blank", description="Template name"),
    output: str = Query(..., description="Output directory path"),
):
    """Scaffold a new project from a domain template."""
    source = TEMPLATES_DIR / template
    if not source.exists():
        available = [d.name for d in TEMPLATES_DIR.iterdir() if d.is_dir()]
        raise HTTPException(404, f"Template '{template}' not found. Available: {available}")

    dest = Path(output).resolve()
    dest.mkdir(parents=True, exist_ok=True)
    shutil.copytree(source, dest, dirs_exist_ok=True)

    return {
        "status": "created",
        "template": template,
        "path": str(dest),
        "project_yml": str(dest / "project.yml"),
    }


@app.post("/project/create")
async def create_project(payload: dict[str, Any] = Body(...)):
    """
    Create a project from the wizard payload.

    Combines template scaffolding with structured edits to ``project.yml``
    so the YAML reflects exactly what the user entered in the wizard
    (no leaked template defaults). Required keys:

      - template: str
      - output:   str (absolute path or relative dir name)
      - identity: {name, description?, author?, response_units?}
      - crs:      {input, projected}
    """
    import yaml

    template = payload.get("template", "blank")
    output = payload.get("output")
    identity = payload.get("identity") or {}
    crs = payload.get("crs") or {}

    if not output:
        raise HTTPException(400, "output directory is required")
    if not identity.get("name"):
        raise HTTPException(400, "identity.name is required")
    if not crs.get("input") or not crs.get("projected"):
        raise HTTPException(400, "crs.input and crs.projected are required")

    source = TEMPLATES_DIR / template
    if not source.exists():
        available = [d.name for d in TEMPLATES_DIR.iterdir() if d.is_dir()]
        raise HTTPException(404, f"Template '{template}' not found. Available: {available}")

    dest = Path(output).resolve()
    dest.mkdir(parents=True, exist_ok=True)
    shutil.copytree(source, dest, dirs_exist_ok=True)

    yml_path = dest / "project.yml"
    if yml_path.exists():
        with open(yml_path, "r", encoding="utf-8") as fh:
            cfg = yaml.safe_load(fh) or {}
    else:
        cfg = {}

    # Merge wizard identity into project block (preserve any template-seeded keys)
    proj = cfg.get("project") or {}
    proj["name"] = identity["name"]
    if "description" in identity:
        proj["description"] = identity.get("description") or ""
    proj["domain"] = template
    if "author" in identity:
        proj["author"] = identity.get("author") or ""
    if "response_units" in identity:
        proj["response_units"] = identity.get("response_units") or ""
    cfg["project"] = proj

    # Merge CRS
    cfg_crs = cfg.get("crs") or {}
    cfg_crs["input"] = crs["input"]
    cfg_crs["projected"] = crs["projected"]
    cfg["crs"] = cfg_crs

    with open(yml_path, "w", encoding="utf-8") as fh:
        yaml.dump(cfg, fh, default_flow_style=False, sort_keys=False)

    return {
        "status": "created",
        "template": template,
        "path": str(dest),
        "project_yml": str(yml_path),
    }


@app.get("/project/templates")
async def list_templates():
    """List available domain templates."""
    templates = []
    for d in sorted(TEMPLATES_DIR.iterdir()):
        if d.is_dir():
            yml = d / "project.yml"
            templates.append({
                "name": d.name,
                "has_project_yml": yml.exists(),
            })
    return {"templates": templates}


@app.get("/project/config")
async def get_project_config():
    """Return the current project configuration as JSON (project.yml structure)."""
    if state.raw_project_yaml is None:
        raise HTTPException(400, "No project loaded.")
    return state.raw_project_yaml


@app.put("/project/config")
async def update_project_config(body: dict[str, Any]):
    """Update the project configuration and persist to disk."""
    if state.raw_project_yaml is None:
        raise HTTPException(400, "No project loaded.")
    state.raw_project_yaml.update(body)
    # Persist to the project YAML
    if state.project_path:
        import yaml
        yml_path = Path(state.project_path)
        if yml_path.is_dir():
            yml_path = yml_path / "project.yml"
        with open(yml_path, "w", encoding="utf-8") as fh:
            yaml.dump(state.raw_project_yaml, fh, default_flow_style=False, sort_keys=False)
        # Reload internal config
        from sparc.config.config import load_config
        state.project_config = load_config(str(yml_path))
    return {"status": "updated"}


# ------------------------------------------------------------------
# Data endpoints
# ------------------------------------------------------------------

@app.get("/data/summary")
async def data_summary():
    """Column statistics for LLM context injection and UI display."""
    if state.data is None:
        raise HTTPException(400, "No data loaded. Load a project first.")
    return state.data_summary or _compute_summary(state)


@app.get("/data/preview")
async def data_preview(n: int = Query(50, ge=1, le=500)):
    """Return the first N rows as JSON records."""
    if state.data is None:
        raise HTTPException(400, "No data loaded.")
    import pandas as pd

    df = state.data.head(n)
    # Drop geometry column for JSON serialization
    if hasattr(df, "geometry"):
        df = pd.DataFrame(df.drop(columns="geometry"))
    return {"rows": df.to_dict(orient="records"), "total": len(state.data)}


@app.get("/data/histogram")
async def data_histogram(
    variable: str = Query(..., description="Column name"),
    bins: int = Query(40, ge=4, le=200),
):
    """
    Server-side histogram for any numeric column. Single pass via numpy;
    returns ~`bins` ints + edges regardless of dataset size. No client-side
    downsampling needed.
    """
    if state.data is None:
        raise HTTPException(400, "No data loaded.")
    if variable not in state.data.columns:
        raise HTTPException(404, f"Column '{variable}' not found")

    import numpy as np
    import pandas as pd

    col = state.data[variable]
    if not pd.api.types.is_numeric_dtype(col):
        raise HTTPException(400, f"Column '{variable}' is not numeric")

    arr = col.to_numpy(dtype="float64", copy=False)
    finite = arr[np.isfinite(arr)]
    n_total = int(arr.size)
    n_finite = int(finite.size)
    n_missing = n_total - n_finite

    if n_finite == 0:
        return {
            "variable": variable,
            "bins": [],
            "edges": [],
            "n_total": n_total,
            "n_finite": 0,
            "n_missing": n_missing,
            "min": None,
            "max": None,
            "mean": None,
            "std": None,
        }

    counts, edges = np.histogram(finite, bins=bins)
    return {
        "variable": variable,
        "bins": counts.tolist(),
        "edges": edges.tolist(),
        "n_total": n_total,
        "n_finite": n_finite,
        "n_missing": n_missing,
        "min": float(finite.min()),
        "max": float(finite.max()),
        "mean": float(finite.mean()),
        "std": float(finite.std(ddof=1)) if n_finite > 1 else 0.0,
    }


@app.get("/crs/distortion")
async def crs_distortion(
    input_epsg: str = Query("4326"),
    projected_epsg: str = Query(""),
):
    """Compute approximate area distortion at the study-area center using pyproj.

    Returns linear scale factor k at the centroid plus estimated area distortion (%).
    This endpoint does NOT require running the pipeline.
    """
    if not projected_epsg:
        raise HTTPException(400, "projected_epsg is required")

    try:
        from pyproj import Transformer, CRS
        import numpy as np

        # Determine study-area center from loaded data or fall back to equator/PM
        cx, cy = 0.0, 45.0  # defaults (lon, lat WGS84)
        if state.data is not None and hasattr(state.data, "geometry"):
            try:
                import geopandas as gpd
                gdf = state.data
                if gdf.crs is not None and str(gdf.crs) != "EPSG:4326":
                    gdf = gdf.to_crs(epsg=4326)
                bounds = gdf.total_bounds  # minx, miny, maxx, maxy
                cx = float((bounds[0] + bounds[2]) / 2)
                cy = float((bounds[1] + bounds[3]) / 2)
            except Exception:
                pass
        elif state.data_summary and "bbox" in (state.data_summary or {}):
            bb = state.data_summary["bbox"]
            if isinstance(bb, dict):
                cx = (bb["minx"] + bb["maxx"]) / 2
                cy = (bb["miny"] + bb["maxy"]) / 2

        # Short-circuit: same CRS → no distortion
        norm_in = input_epsg.replace("EPSG:", "").strip()
        norm_proj = projected_epsg.replace("EPSG:", "").strip()
        if norm_in == norm_proj:
            src_crs = CRS.from_epsg(int(norm_in))
            return {
                "center_lon": cx,
                "center_lat": cy,
                "k_x": 1.0,
                "k_y": 1.0,
                "k_mean": 1.0,
                "area_distortion_pct": 0.0,
                "input_crs_name": src_crs.name,
                "projected_crs_name": src_crs.name,
                "assessment": "acceptable",
            }

        # Compute k (linear scale factor) numerically: transform two points 1 m apart
        # in WGS84, see how far they are in projected CRS vs expected
        delta_deg = 0.001  # ~100m along latitude
        try:
            t = Transformer.from_crs(
                f"EPSG:{input_epsg.replace('EPSG:', '')}",
                f"EPSG:{projected_epsg.replace('EPSG:', '')}",
                always_xy=True,
            )
            x0, y0 = t.transform(cx, cy)
            x1, y1 = t.transform(cx + delta_deg, cy)
            x2, y2 = t.transform(cx, cy + delta_deg)

            # Approximate ellipsoidal distance (metres) for the same delta_deg
            from pyproj import Geod
            geod = Geod(ellps="WGS84")
            _, _, dist_x = geod.inv(cx, cy, cx + delta_deg, cy)
            _, _, dist_y = geod.inv(cx, cy, cx, cy + delta_deg)

            # If projected CRS has angular units (i.e. it is also geographic),
            # convert projected delta from degrees to metres using geod distances.
            tgt_crs_obj = CRS.from_epsg(int(norm_proj))
            tgt_is_angular = tgt_crs_obj.axis_info[0].unit_name in ("degree", "grad")
            if tgt_is_angular:
                # Both input and output are in degrees — distances are the same
                k_x = 1.0
                k_y = 1.0
            else:
                k_x = float(np.sqrt((x1 - x0) ** 2 + (y1 - y0) ** 2) / dist_x) if dist_x else 1.0
                k_y = float(np.sqrt((x2 - x0) ** 2 + (y2 - y0) ** 2) / dist_y) if dist_y else 1.0
            k_mean = (k_x + k_y) / 2
            area_distortion_pct = float(abs(k_mean ** 2 - 1) * 100)

            src_crs = CRS.from_epsg(int(input_epsg.replace("EPSG:", "")))
            tgt_crs = CRS.from_epsg(int(projected_epsg.replace("EPSG:", "")))

            return {
                "center_lon": cx,
                "center_lat": cy,
                "k_x": k_x,
                "k_y": k_y,
                "k_mean": k_mean,
                "area_distortion_pct": area_distortion_pct,
                "input_crs_name": src_crs.name,
                "projected_crs_name": tgt_crs.name,
                "assessment": "acceptable" if area_distortion_pct < 0.5 else "high",
            }
        except Exception as e:
            raise HTTPException(400, f"Projection error: {e}")

    except ImportError:
        raise HTTPException(500, "pyproj not available")


@app.get("/data/geojson")
async def data_geojson(variable: str | None = Query(None)):
    """Return raw data as GeoJSON, optionally filtered to a single variable for map coloring."""
    if state.data is None:
        raise HTTPException(400, "No data loaded.")

    df = state.data
    if not hasattr(df, "geometry") or df.geometry is None:
        raise HTTPException(400, "Loaded data has no geometry column.")

    if variable:
        if variable not in df.columns:
            raise HTTPException(400, f"Column '{variable}' not found.")
        subset = df[["geometry", variable]].copy()
    else:
        # Return just geometry + first 5 numeric cols to keep payload small
        numeric_cols = list(df.select_dtypes(include="number").columns[:5])
        subset = df[["geometry"] + numeric_cols].copy()

    # Reproject to WGS84 for web display (deck.gl / maplibre expect EPSG:4326)
    if hasattr(subset, "crs") and subset.crs is not None and str(subset.crs) != "EPSG:4326":
        subset = subset.to_crs(epsg=4326)

    return subset.__geo_interface__


@app.post("/data/upload")
async def upload_data(file: UploadFile = File(...)):
    """Accept a CSV, raster (.tif/.tiff), or spatial file (.shp/.gpkg/.geojson) upload."""
    if state.project_config is None:
        raise HTTPException(400, "Load a project first.")

    project_dir = Path(state.project_config["paths"]["project_root"])
    data_dir = project_dir / "data"
    data_dir.mkdir(parents=True, exist_ok=True)

    safe_name = Path(file.filename).name  # strip any directory components
    dest = data_dir / safe_name
    content = await file.read()
    with open(dest, "wb") as f:
        f.write(content)

    suffix = dest.suffix.lower()
    result: dict = {"status": "uploaded", "path": str(dest), "file_type": suffix}

    # CSV / Parquet — load as primary data
    if suffix in (".csv", ".parquet"):
        _set_data_path(str(dest))
        _load_data_into_state(state.project_config)
        result.update({
            "columns": list(state.data.columns) if state.data is not None else [],
            "row_count": len(state.data) if state.data is not None else 0,
        })

    # Raster — store metadata for zonal-stats pipeline
    elif suffix in (".tif", ".tiff"):
        try:
            import rasterio
            with rasterio.open(dest) as src:
                result.update({
                    "crs": str(src.crs) if src.crs else None,
                    "bounds": list(src.bounds),
                    "shape": list(src.shape),
                    "band_count": src.count,
                })
        except ImportError:
            result["crs"] = None
            result["note"] = "rasterio not installed; metadata unavailable"
        except Exception as exc:
            result["note"] = f"Could not read raster metadata: {exc}"

    # Shapefile / GeoPackage / GeoJSON — boundary or vector layer
    elif suffix in (".shp", ".gpkg", ".geojson"):
        try:
            import geopandas as _gpd
            gdf = _gpd.read_file(dest)
            result.update({
                "crs": str(gdf.crs) if gdf.crs else None,
                "n_features": len(gdf),
                "columns": [c for c in gdf.columns if c != "geometry"],
                "bounds": list(gdf.total_bounds),
            })
        except Exception as exc:
            result["note"] = f"Could not read spatial file: {exc}"

    # Shapefile sidecar files (.shx, .dbf, .prj, .cpg) — just store
    elif suffix in (".shx", ".dbf", ".prj", ".cpg"):
        result["note"] = "Shapefile sidecar stored"

    else:
        result["note"] = f"Unknown file type {suffix}; stored as-is"

    return result


@app.get("/data/files")
async def list_data_files():
    """List CSV/Parquet files inside the project's working directory."""
    if state.project_config is None:
        raise HTTPException(400, "Load a project first.")

    project_dir = Path(state.project_config["paths"]["project_root"])
    files: list[dict] = []
    for ext in ("*.csv", "*.parquet", "*.CSV", "*.tif", "*.tiff", "*.shp", "*.gpkg", "*.geojson"):
        for p in project_dir.rglob(ext):
            try:
                rel = p.relative_to(project_dir)
            except ValueError:
                rel = p
            files.append({
                "name": p.name,
                "path": str(p),
                "relative": str(rel),
                "size": p.stat().st_size,
            })
    files.sort(key=lambda f: f["relative"])
    return {"project_dir": str(project_dir), "files": files}


@app.post("/data/select")
async def select_data_file(path: str = Query(..., description="Absolute path to a data file")):
    """Select an existing data file (must already be on disk)."""
    if state.project_config is None:
        raise HTTPException(400, "Load a project first.")

    resolved = Path(path).resolve()
    if not resolved.exists():
        raise HTTPException(404, f"File not found: {resolved}")
    if resolved.suffix.lower() not in (".csv", ".parquet"):
        raise HTTPException(400, "Only .csv and .parquet files are supported.")

    # Update config to point to this file and persist to disk
    _set_data_path(str(resolved))
    _load_data_into_state(state.project_config)

    return {
        "status": "selected",
        "path": str(resolved),
        "columns": list(state.data.columns) if state.data is not None else [],
        "row_count": len(state.data) if state.data is not None else 0,
    }


# ------------------------------------------------------------------
# Data validation endpoint
# ------------------------------------------------------------------

@app.post("/data/validate")
async def validate_data():
    """Run pre-flight validation checks on the currently loaded dataset.

    Returns a structured checklist of issues (missing values, CRS,
    duplicates, distribution anomalies, type errors) with severity
    levels: critical / warning / info.
    """
    if state.data is None:
        raise HTTPException(400, "No data loaded. Load a project first.")
    if state.project_config is None:
        raise HTTPException(400, "No project loaded.")

    from sparc.data.validation import validate_dataset

    config = state.project_config
    report = validate_dataset(
        state.data,
        target_col=config.get("variables", {}).get("target"),
        predictor_cols=config.get("predictors", {}).get("base_model", []),
        coord_cols=config.get("variables", {}).get("coordinates", []),
        expected_crs=config.get("crs", {}).get("initial"),
    )
    return report.to_dict()


# ------------------------------------------------------------------
# Data versioning endpoints
# ------------------------------------------------------------------

@app.get("/data/versions")
async def list_data_versions():
    """Return all versioned data snapshots for the current project."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded.")

    from sparc.data.versioning import list_versions

    project_dir = Path(state.project_config["paths"]["project_root"])
    data_dir = project_dir / "data"
    versions = list_versions(data_dir)
    return {"versions": versions}


@app.post("/data/select_version")
async def select_data_version(version: int = Query(..., description="Version number to activate")):
    """Switch the active dataset to a specific versioned snapshot."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded.")

    from sparc.data.versioning import get_version_path

    project_dir = Path(state.project_config["paths"]["project_root"])
    data_dir = project_dir / "data"
    path = get_version_path(data_dir, version)

    if path is None:
        raise HTTPException(404, f"Version {version} not found or file missing.")

    _set_data_path(str(path))
    _load_data_into_state(state.project_config)

    return {
        "status": "selected",
        "version": version,
        "path": str(path),
        "columns": list(state.data.columns) if state.data is not None else [],
        "row_count": len(state.data) if state.data is not None else 0,
    }


# ------------------------------------------------------------------
# Session log endpoint
# ------------------------------------------------------------------

@app.get("/run/log")
async def get_run_log():
    """Return the persistent session log for the current project.

    Each entry is a JSON object with timestamp, stage, type, message, and data.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded.")

    project_dir = Path(state.project_config["paths"]["project_root"])
    log_path = project_dir / "session_log.jsonl"

    if not log_path.exists():
        return {"entries": [], "path": str(log_path)}

    import json
    entries = []
    with open(log_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue

    return {"entries": entries, "path": str(log_path)}


# ------------------------------------------------------------------
# Pipeline streaming (WebSocket)
# ------------------------------------------------------------------

@app.websocket("/run/stream")
async def run_stream(ws: WebSocket):
    """Stream structured pipeline events over a WebSocket.

    The client sends a JSON message to start:
        ``{"stage": 2, "fast": false, "skip_gwen": false}``
        ``{"subscribe": "artifacts"}``  — no stage executed; receive artifact events only

    The server pushes events until the stage completes or errors:
        ``{"type": "metric", "stage": 2, "fold": 3, "metric": "r2", "value": 0.891}``
        ``{"type": "complete", "stage": 2}``
        ``{"type": "artifact_written", "stage": "4", "artifact_id": "scenario_results", ...}``

    ``artifact_written`` events are emitted whenever ``RunRegistry.register_artifact``
    is called, so the desktop reacts the moment new data lands in the database.
    """
    await ws.accept()

    try:
        init_msg = await ws.receive_json()
    except WebSocketDisconnect:
        return

    # Subscribe-only mode: client wants artifact events but is not starting a run.
    if init_msg.get("subscribe") == "artifacts" or init_msg.get("stage") is None:
        queue: asyncio.Queue = asyncio.Queue()
        queue._sparc_loop = asyncio.get_running_loop()  # type: ignore[attr-defined]
        _artifact_subscribers.append(queue)
        # Ack so clients (and tests) know the subscription is live before
        # they trigger downstream writes.
        await ws.send_json({"type": "subscribed", "channel": "artifacts"})
        try:
            while True:
                event = await queue.get()
                await ws.send_json(event)
        except WebSocketDisconnect:
            pass
        finally:
            try:
                _artifact_subscribers.remove(queue)
            except ValueError:
                pass
            if ws.client_state.name != "DISCONNECTED":
                await ws.close()
        return

    stage = int(init_msg.get("stage", 0))
    fast = bool(init_msg.get("fast", False))
    skip_gwen = bool(init_msg.get("skip_gwen", False))

    if state.project_config is None:
        await ws.send_json({"type": "error", "message": "No project loaded"})
        await ws.close()
        return

    if state.is_running:
        await ws.send_json({"type": "error", "message": f"Stage {state.current_stage} already running"})
        await ws.close()
        return

    # Subscribe to artifact events while running so the same socket carries
    # both metric/complete and artifact_written events.
    queue = asyncio.Queue()
    queue._sparc_loop = asyncio.get_running_loop()  # type: ignore[attr-defined]
    _artifact_subscribers.append(queue)

    async def _drain_artifacts() -> None:
        try:
            while True:
                event = await queue.get()
                await ws.send_json(event)
        except (asyncio.CancelledError, WebSocketDisconnect):
            pass

    drain_task = asyncio.create_task(_drain_artifacts())
    try:
        async for event in stream_stage(state, stage, fast=fast, skip_gwen=skip_gwen):
            await ws.send_json(event)
    except WebSocketDisconnect:
        pass
    finally:
        drain_task.cancel()
        try:
            _artifact_subscribers.remove(queue)
        except ValueError:
            pass
        if ws.client_state.name != "DISCONNECTED":
            await ws.close()


@app.get("/run/events")
async def get_run_events():
    """Return buffered events for the current (or most recent) pipeline run.

    Allows a client that reconnects (e.g. navigated away) to catch up on
    events it missed without needing to re-open the WebSocket.
    """
    return {
        "is_running": state.is_running,
        "current_stage": state.current_stage,
        "events": state.get_buffered_events(),
    }


# ------------------------------------------------------------------
# Run comparison (Phase 15) — diff two output directories
# ------------------------------------------------------------------

@app.get("/runs/discover")
async def get_runs_discover(search_root: str | None = Query(default=None)):
    """List candidate output directories the user could compare.

    If *search_root* is omitted we default to the project's ``output_dir``
    parent (so siblings of the current run show up). Falls back to the
    user's home directory.
    """
    from sparc.evaluation.diff import discover_runs

    if search_root:
        root = Path(search_root)
    else:
        root = Path.home()
        if state.project_config is not None:
            try:
                from sparc.run.pipeline_paths import PipelinePaths
                paths = PipelinePaths.from_config(state.project_config)
                if paths.output_dir.parent.exists():
                    root = paths.output_dir.parent
            except Exception:
                pass
    return {"search_root": str(root), "runs": discover_runs(root)}


@app.get("/runs/diff")
async def get_runs_diff(
    run_a: str = Query(...),
    run_b: str = Query(...),
):
    """Diff two pipeline output directories on disk."""
    from sparc.evaluation.diff import diff_runs

    try:
        return diff_runs(Path(run_a), Path(run_b))
    except FileNotFoundError as exc:
        raise HTTPException(404, str(exc))
    except Exception as exc:
        raise HTTPException(500, f"Diff failed: {exc}")


# ------------------------------------------------------------------
# Reproducibility (Phase 17)
# ------------------------------------------------------------------

@app.get("/reproduce/provenance")
async def get_provenance(run_dir: str = Query(...)):
    """Return ``run_provenance.json`` for the given run directory (or 404)."""
    from sparc.registry.provenance import load_provenance, load_frozen_config
    p = load_provenance(run_dir)
    if p is None:
        raise HTTPException(404, f"No run_provenance.json under {run_dir}")
    cfg = load_frozen_config(run_dir)
    return {"provenance": p, "frozen_config_present": cfg is not None}


@app.post("/reproduce/freeze")
async def post_reproduce_freeze():
    """Freeze the *current* project: write run_provenance.json + frozen_config.json
    to the active output_dir.
    """
    from sparc.registry.provenance import freeze_run
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    try:
        from sparc.run.pipeline_paths import PipelinePaths
        paths = PipelinePaths.from_config(state.project_config)
        out = paths.output_dir
    except Exception as exc:
        raise HTTPException(500, f"Cannot resolve output_dir: {exc}")
    repo_root = Path(__file__).resolve().parents[2]
    p = freeze_run(out, state.project_config, repo_root=repo_root)
    return {"path": str(p), "output_dir": str(out)}


@app.post("/reproduce/load")
async def post_reproduce_load(payload: dict = Body(...)):
    """Replace the active project_config with the frozen config from a run dir.

    The client should then trigger ``/run/stream`` as usual to actually
    execute the pipeline. After it completes, call ``/reproduce/verify``
    to compare provenance against the original run.
    """
    from sparc.registry.provenance import load_frozen_config, load_provenance
    run_dir = payload.get("run_dir")
    if not run_dir:
        raise HTTPException(400, "run_dir required")
    cfg = load_frozen_config(run_dir)
    if cfg is None:
        raise HTTPException(404, f"frozen_config.json not found in {run_dir}")
    state.project_config = cfg
    prov = load_provenance(run_dir)
    return {
        "loaded": True,
        "config_hash": (prov or {}).get("config_hash"),
        "data": (prov or {}).get("data"),
        "warnings": ["Re-run via /run/stream to reproduce. Then call /reproduce/verify."],
    }


@app.get("/reproduce/verify")
async def get_reproduce_verify(
    run_a: str = Query(...),
    run_b: str = Query(...),
):
    """Compare provenance between two runs (config/data/git/env)."""
    from sparc.registry.provenance import compare_runs, verify_sidecars
    cmp = compare_runs(run_a, run_b)
    cmp["sidecars_b"] = verify_sidecars(run_b)
    return cmp


# ------------------------------------------------------------------
# Scenario library (Phase 17)
# ------------------------------------------------------------------

def _scenario_library_dir() -> Path:
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    from sparc.run.pipeline_paths import PipelinePaths
    paths = PipelinePaths.from_config(state.project_config)
    return paths.output_dir


@app.get("/scenarios/library")
async def get_scenario_library():
    from sparc.scenario.library import build_timeline
    return build_timeline(_scenario_library_dir())


@app.post("/scenarios/library")
async def post_scenario_library(payload: dict = Body(...)):
    from sparc.scenario.library import append_entry
    from sparc.registry.provenance import compute_config_hash
    scenario = payload.get("scenario")
    if not isinstance(scenario, dict):
        raise HTTPException(400, "scenario (dict) required")
    cfg_hash = compute_config_hash(state.project_config or {}) if state.project_config else None
    entry = append_entry(
        _scenario_library_dir(),
        scenario,
        author=str(payload.get("author") or "anonymous"),
        comment=str(payload.get("comment") or ""),
        parent_id=payload.get("parent_id"),
        config_hash=cfg_hash,
    )
    return entry


# ------------------------------------------------------------------
# Standalone snapshot HTML (Phase 17)
# ------------------------------------------------------------------

@app.post("/report/standalone")
async def post_report_standalone(payload: dict | None = Body(default=None)):
    """Bundle the active run into a single self-contained HTML.

    Body (all optional):
        {
          "run_dir": "<path>",       # defaults to active project's output_dir
          "chat_history": [{role, content}, ...],
          "project_name": "..."
        }
    """
    from sparc.report.standalone_html import build_standalone_html
    from starlette.responses import Response as StarletteResponse

    payload = payload or {}
    run_dir = payload.get("run_dir")
    if not run_dir:
        if state.project_config is None:
            raise HTTPException(400, "No run_dir and no project loaded")
        from sparc.run.pipeline_paths import PipelinePaths
        run_dir = str(PipelinePaths.from_config(state.project_config).output_dir)
    project_name = payload.get("project_name")
    if not project_name and state.project_config is not None:
        project_name = (state.project_config.get("project") or {}).get("name")
    chat_history = payload.get("chat_history") if isinstance(payload.get("chat_history"), list) else None
    html_text = build_standalone_html(run_dir, chat_history=chat_history, project_name=project_name)
    fname = f"sparc_snapshot_{Path(run_dir).name or 'run'}.html"
    return StarletteResponse(
        content=html_text,
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ------------------------------------------------------------------
# Context layers (Phase 18)
# ------------------------------------------------------------------

@app.get("/context/layers")
async def get_context_layers(domain: str | None = Query(default=None)):
    """Return the basemap+overlay tile catalog, optionally filtered by *domain*.

    If *domain* is omitted we infer it from the active project's
    ``project.domain`` so the UI can recommend defaults without an extra
    round-trip.
    """
    from sparc.server.context_layers import get_catalog
    if domain is None and state.project_config is not None:
        proj = state.project_config.get("project") or {}
        domain = proj.get("domain")
    return get_catalog(domain)


# ------------------------------------------------------------------
# Audience-specific reports (Phase 19)
# ------------------------------------------------------------------

@app.post("/report/audience")
async def post_report_audience(
    audience: str = Query(..., regex="^(technical|planner|public)$"),
    fmt: str = Query("html", regex="^(md|html|pdf)$"),
    payload: dict | None = Body(default=None),
):
    """Render a SPARC run for a specific audience.

    Query params:
      audience: technical | planner | public
      fmt:      md | html | pdf

    Body (optional):
      { "run_dir": "<path>" }   # defaults to active project's output_dir

    Returns a file download (Content-Disposition attachment).
    """
    from sparc.report.audience import generate_audience_report
    from starlette.responses import Response as StarletteResponse

    payload = payload or {}
    run_dir = payload.get("run_dir")
    if not run_dir:
        if state.project_config is None:
            raise HTTPException(400, "No project loaded and no run_dir provided")
        run_dir = state.project_config.get("paths", {}).get("output_dir")
    if not run_dir:
        raise HTTPException(400, "Cannot resolve run_dir")
    run_path = Path(run_dir)
    if not run_path.exists():
        raise HTTPException(404, f"Run directory not found: {run_dir}")

    try:
        result = generate_audience_report(
            run_dir=run_path,
            config=state.project_config,
            audience=audience,
            fmt=fmt,
        )
    except RuntimeError as exc:
        raise HTTPException(501, str(exc))
    except ValueError as exc:
        raise HTTPException(400, str(exc))

    media = {"md": "text/markdown", "html": "text/html", "pdf": "application/pdf"}[fmt]
    fname = f"sparc_{audience}.{fmt}"
    content = result if isinstance(result, (bytes, bytearray)) else result.encode("utf-8")
    return StarletteResponse(
        content=content,
        media_type=media,
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


# ------------------------------------------------------------------
# Structured results endpoints (MUST be defined before /results/{stage}
# so FastAPI matches exact paths before the parameterized catch-all)
# ------------------------------------------------------------------

@app.get("/results/correlogram")
async def get_correlogram_data():
    """Return correlogram analysis results with per-variable lag/Moran's I data."""
    return _read_or_404(
        "0", "correlogram_results",
        hint="Stage 0 (EDA) has not produced correlogram_results. Run Stage 0.",
    )


@app.get("/results/gwen")
async def get_gwen_data():
    """Return GWEN variable importance as a row-oriented table."""
    store = _open_store()
    try:
        if store.has("1", "gwen_variable_importance"):
            df = store.read_any("1", "gwen_variable_importance")
            return {"rows": df.to_dict(orient="records")}
        if store.has("1", "gwen_results"):
            return store.read_any("1", "gwen_results")
    finally:
        from sparc.registry.run_registry import set_active_registry
        set_active_registry(None)
    raise _missing_artifact_response(
        artifact_id="gwen_variable_importance", stage="1",
        hint="Stage 1 (GWEN) has not produced gwen_variable_importance. Run Stage 1.",
    )


@app.get("/results/model_performance")
async def get_model_performance():
    """Return per-model R²/RMSE for the R² bar chart on the Results page."""
    models: list[dict] = []

    # In-memory result from the just-finished stage 2 (if available).
    r2_result = state.get_result(2)
    if isinstance(r2_result, dict):
        indiv = (r2_result.get("performance") or {}).get("individual_models")
        if indiv is None:
            indiv = r2_result.get("base_models")
        if isinstance(indiv, dict):
            for name, metrics in indiv.items():
                if isinstance(metrics, dict) and metrics.get("r2") is not None:
                    models.append({
                        "name": name.upper(),
                        "r2": metrics["r2"],
                        "rmse": metrics.get("rmse"),
                    })
        ens = r2_result.get("final_ensemble") or r2_result.get("meta_ensemble_best")
        if isinstance(ens, dict) and ens.get("r2") is not None:
            models.append({
                "name": "Ensemble",
                "r2": ens["r2"],
                "rmse": ens.get("rmse"),
            })

    if not models:
        store = _open_store()
        try:
            if store.has("2", "ensemble_results"):
                data = store.read_any("2", "ensemble_results")
                base = data.get("base_models", {})
                for name, metrics in base.items():
                    if isinstance(metrics, dict) and metrics.get("r2") is not None:
                        models.append({
                            "name": name.upper(),
                            "r2": metrics["r2"],
                            "rmse": metrics.get("rmse"),
                        })
                ens = data.get("final_ensemble") or data.get("meta_ensemble_best")
                if isinstance(ens, dict) and ens.get("r2") is not None:
                    models.append({
                        "name": "Ensemble",
                        "r2": ens["r2"],
                        "rmse": ens.get("rmse"),
                    })
        finally:
            from sparc.registry.run_registry import set_active_registry
            set_active_registry(None)

    if not models:
        raise _missing_artifact_response(
            artifact_id="ensemble_results", stage="2",
            hint="Stage 2 has not produced ensemble_results. Run Stage 2.",
        )

    models.sort(key=lambda m: m.get("r2") or 0, reverse=True)
    return {"models": models}


@app.get("/results/spatial_cv/predictions")
async def get_spatial_cv_predictions():
    """Return spatial CV predictions as GeoJSON (DB-only)."""
    store = _open_store()
    try:
        if not store.has("2", "spatial_cv_predictions"):
            raise _missing_artifact_response(
                artifact_id="spatial_cv_predictions", stage="2",
                hint=(
                    "Stage 2 has not produced spatial_cv_predictions. "
                    "Run Stage 2 to populate the predictions map."
                ),
            )
        gdf = store.read_any("2", "spatial_cv_predictions")
    finally:
        from sparc.registry.run_registry import set_active_registry
        set_active_registry(None)

    import geopandas as gpd
    if not isinstance(gdf, gpd.GeoDataFrame):
        raise HTTPException(
            500,
            "Artifact (2, spatial_cv_predictions) is not a GeoDataFrame; "
            "the producer must include a geometry column.",
        )
    if gdf.crs is not None and str(gdf.crs) != "EPSG:4326":
        gdf = gdf.to_crs(epsg=4326)
    return gdf.__geo_interface__


@app.get("/results/causal")
async def get_causal_results():
    """Return causal validation results (DB-only)."""
    store = _open_store()
    try:
        if store.has("3", "scenario_coefficients"):
            return store.read_any("3", "scenario_coefficients")
        if store.has("3", "causal_diagnostics"):
            return store.read_any("3", "causal_diagnostics")
    finally:
        from sparc.registry.run_registry import set_active_registry
        set_active_registry(None)

    mem_result = state.get_result(3)
    if mem_result is not None:
        return mem_result

    raise _missing_artifact_response(
        artifact_id="scenario_coefficients", stage="3",
        hint=(
            "Stage 3 (Causal Validation) has not produced "
            "`scenario_coefficients`. Run Stage 3 to populate the causal panel."
        ),
    )


@app.get("/results/causal/dose_response")
async def get_dose_response():
    """Return dose-response curves (DB-only)."""
    return _read_or_404(
        "3", "dose_response_curves",
        hint=(
            "Dose-response curves are produced by Stage 3 (Causal Validation). "
            "They are skipped when `causal.inference: bayesian`. Re-run Stage 3 "
            "in frequentist mode or check the Causal Diagnostics panel for the "
            "Bayesian posterior."
        ),
    )


@app.get("/results/causal/sensitivity")
async def get_causal_sensitivity():
    """Return E-values + tipping-point analysis for the current causal payload.

    Reuses ``/results/causal`` and annotates each effect with VanderWeele
    E-values so users can see how strong an unmeasured confounder would
    have to be to explain away the observed effect.
    """
    from sparc.causal.sensitivity import annotate_causal_payload

    payload = await get_causal_results()  # type: ignore[misc]
    if not isinstance(payload, dict):
        raise HTTPException(500, "Causal payload is not a dict")
    annotated = annotate_causal_payload(payload)
    sens = annotated.get("sensitivity") or {}
    return sens


@app.get("/results/causal/negative_control")
async def get_causal_negative_control(
    variable: str = Query(...),
    n_permutations: int = Query(1000, ge=50, le=10000),
):
    """Permutation negative-control test on the spatial CATE values (DB-only)."""
    from sparc.evaluation.negative_controls import permutation_test_cate

    arr_np = _read_cate_multiplier_from_store(variable)
    if arr_np is None:
        raise _missing_artifact_response(
            artifact_id=f"cate_summary[{variable}]", stage="3",
            hint=(
                f"No CATE map for variable '{variable}'. Re-run Stage 3 "
                "with `causal.estimate_cate: true`."
            ),
        )
    arr = arr_np.astype(float).tolist()
    res = permutation_test_cate(arr, n_permutations=n_permutations)
    return {
        "variable": variable,
        "n": res.n,
        "mean_observed": res.mean_observed,
        "mean_null": res.mean_null,
        "std_null": res.std_null,
        "p_value": res.p_value,
        "z_score": res.z_score,
        "n_permutations": res.n_permutations,
        "passed": res.passed,
        "interpretation": (
            "p > 0.05 — observed CATE indistinguishable from null (negative-control PASS)."
            if res.passed
            else f"p = {res.p_value:.4f} — CATE is significantly non-zero (real treatment effect)."
        ),
    }


@app.get("/results/causal/diagnostics")
async def get_causal_diagnostics():
    """Return CATE diagnostics (calibration, cumulative effects, RATE)."""
    return _read_or_404(
        "3", "causal_diagnostics",
        hint="Stage 3 (Causal Validation) has not produced causal_diagnostics.",
    )


def _load_neural_pdp(paths) -> dict:
    """Read PINN neural-network PDP curves into the canonical curves dict (DB-only).

    Each PDP table is registered as ``v2_neural_pdp::<feature>`` under stage 2.
    The ``paths`` argument is accepted for API compatibility but unused.
    """
    out: dict = {}
    if state.registry is None:
        return out
    try:
        from sparc.registry.run_registry import set_active_registry
        from sparc.registry.store import ArtifactStore
        set_active_registry(state.registry)
        try:
            _store = ArtifactStore(state.registry)
            manifest_stages = getattr(state.registry.manifest, "stages", {}) or {}
            stage2 = manifest_stages.get("2", {}) or {}
            for art_id in stage2.keys():
                if not art_id.startswith("v2_neural_pdp::"):
                    continue
                feat_col = art_id.split("::", 1)[1]
                try:
                    df = _store.read_any("2", art_id)
                except Exception:
                    continue
                if df is None or len(df) == 0 or feat_col not in df.columns:
                    continue
                grid = [float(v) for v in df[feat_col].tolist()]
                pdp_vals = [float(v) for v in df["mean_prediction"].tolist()]
                if {"q10", "q90"}.issubset(df.columns):
                    pdp_std = [
                        (float(q90) - float(q10)) / 2.56
                        for q10, q90 in zip(df["q10"].tolist(), df["q90"].tolist())
                    ]
                else:
                    pdp_std = None
                out[feat_col] = {
                    "grid_values": grid,
                    "pdp_values": pdp_vals,
                    "pdp_std": pdp_std,
                    "source": "neural_pde",
                }
        finally:
            set_active_registry(None)
    except Exception:
        pass
    return out


def _load_gwrf_pdp(paths) -> dict | None:
    """Load GWRF condition curves dict from artifacts.db (DB-only).

    Returns the parsed dict (variable -> curve metadata) or ``None`` if no
    artifact is registered. The ``paths`` argument is accepted for API
    compatibility but unused.
    """
    if state.registry is None:
        return None
    try:
        from sparc.registry.run_registry import set_active_registry
        from sparc.registry.store import ArtifactStore
        set_active_registry(state.registry)
        try:
            _store = ArtifactStore(state.registry)
            for stage_id, art_id in [
                ("2", "gwrf_condition_curves"),
                ("3", "gwrf_condition_curves"),
            ]:
                if _store.has(stage_id, art_id):
                    data = _store.read_any(stage_id, art_id)
                    if isinstance(data, dict):
                        for _var, curve in data.items():
                            if isinstance(curve, dict):
                                curve.setdefault("source", "gwrf")
                        return data
        finally:
            set_active_registry(None)
    except Exception:
        return None
    return None


@app.get("/results/neural_pdp")
async def get_neural_pdp():
    """Return PINN-derived PDP curves (DB-only)."""
    curves = _load_neural_pdp(None) if state.registry is None else _load_neural_pdp(None)
    if not curves:
        raise _missing_artifact_response(
            artifact_id="v2_neural_pdp::*", stage="2",
            hint="Neural PDP curves not available — run Stage 2 (v2_neural training) first.",
        )
    return curves


@app.get("/results/pdp_curves")
async def get_pdp_curves():
    """Return partial dependence / condition curves from any available source (DB-only).

    Merges in order of preference:
      1. PINN neural-network PDP   (artifact `v2_neural_pdp::<feature>`)
      2. GWRF condition curves     (artifact `gwrf_condition_curves`)
    Causal dose-response (Stage 3) is exposed under `_meta.by_source.causal_dose_response`.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    gwrf = _load_gwrf_pdp(None) or {}
    neural = _load_neural_pdp(None) or {}

    available_sources: list[str] = []
    if neural:
        available_sources.append("neural_pde")
    if gwrf:
        available_sources.append("gwrf")

    causal_curves: dict = {}
    if state.registry is not None:
        try:
            from sparc.registry.run_registry import set_active_registry
            from sparc.registry.store import ArtifactStore
            set_active_registry(state.registry)
            try:
                _store = ArtifactStore(state.registry)
                if _store.has("3", "dose_response_curves"):
                    cd = _store.read_any("3", "dose_response_curves")
                    if isinstance(cd, dict):
                        for var, curve in cd.items():
                            if isinstance(curve, dict):
                                curve.setdefault("source", "causal_dose_response")
                                causal_curves[var] = curve
                        if causal_curves:
                            available_sources.append("causal_dose_response")
            finally:
                set_active_registry(None)
        except Exception:
            pass

    merged: dict = {}
    merged.update(gwrf)
    merged.update(neural)

    if not merged and not causal_curves:
        raise _missing_artifact_response(
            artifact_id="pdp_curves", stage="2",
            hint=(
                "No response-curve data found. Stage 2 produces neural PDP "
                "and GWRF condition curves; Stage 3 (frequentist mode) produces "
                "causal dose-response curves. Re-run the relevant stage."
            ),
        )

    merged["_meta"] = {
        "available_sources": available_sources,
        "by_source": {
            "neural_pde": neural,
            "gwrf": gwrf,
            "causal_dose_response": causal_curves,
        },
    }
    return merged


@app.get("/results/scenarios/nuts_summary")
async def get_nuts_summary():
    """Return NUTS posterior summaries, convergence diagnostics, BMA coefficients (DB-only)."""
    store = _open_store()
    try:
        result: dict = {}
        if store.has("3", "nuts_summary"):
            ns = store.read_any("3", "nuts_summary") or {}
            result["acceptance_rate"] = ns.get("acceptance_rate")
            result["n_divergences"] = ns.get("n_divergences")
        if store.has("3", "parameter_posteriors"):
            df = store.read_any("3", "parameter_posteriors")
            if df is not None:
                result["posteriors"] = df.to_dict(orient="records")
        if store.has("3", "convergence_diagnostics"):
            df = store.read_any("3", "convergence_diagnostics")
            if df is not None:
                result["convergence"] = df.to_dict(orient="records")
        if store.has("3", "bma_coefficients"):
            df = store.read_any("3", "bma_coefficients")
            if df is not None:
                result["bma"] = df.to_dict(orient="records")
    finally:
        from sparc.registry.run_registry import set_active_registry
        set_active_registry(None)

    if not result:
        raise _missing_artifact_response(
            artifact_id="nuts_summary", stage="3",
            hint=(
                "No NUTS results found. Run Stage 3 with "
                "`causal.inference: bayesian` to populate posterior summaries."
            ),
        )
    return result


@app.get("/results/scenarios")
async def list_scenarios():
    """Enumerate available scenarios from the active scenario_results table.

    Reads exclusively through the ArtifactStore (no disk fallback). Returns
    `{ "scenarios": [{"index": int, "pred_column": str, "delta_column": str|None}],
       "results_artifact_id": str | None,
       "available": [str] }`
    """
    from sparc.registry.store import ArtifactStore
    from sparc.scenario import ScenarioBundle

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint="No active run registry. Load a project first.",
        )
    store = ArtifactStore(state.registry)
    bundle = ScenarioBundle.from_store(store)
    return {
        "results_artifact_id": bundle.results_artifact_id,
        "summary_artifact_id": bundle.summary_artifact_id,
        "available": sorted(bundle.available),
        "has_uncertainty": bundle.has_uncertainty(),
        "scenarios": [
            {
                "index": sid,
                "pred_column": (rec := bundle.get_scenario(sid)).pred_column,
                "delta_column": rec.delta_column,
            }
            for sid in bundle.list_scenarios()
        ],
    }


@app.get("/results/scenarios/detail")
async def get_scenario_detail():
    """Return scenario results as GeoJSON with delta columns + summary table.

    Reads exclusively through the ArtifactStore via :class:`ScenarioBundle`.
    Mode-variant precedence (hybrid > reprediction > dag > base) is handled
    automatically. The legacy on-disk ``scenario_results*.gpkg`` files are
    deliberately ignored.
    """
    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint="No active run registry. Load a project first.",
        )

    from sparc.registry.store import ArtifactStore
    from sparc.scenario import (
        DELTA_SCENARIO_PREFIX,
        PRED_BASELINE_COL,
        PRED_SCENARIO_PREFIX,
        ScenarioBundle,
    )

    store = ArtifactStore(state.registry)
    bundle = ScenarioBundle.from_store(store)

    if "results" not in bundle.available or bundle.results is None:
        scenarios_cfg = (state.project_config or {}).get("scenarios") or []
        scenario_count = len(scenarios_cfg) if isinstance(scenarios_cfg, list) else 0
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint=(
                f"Stage 4 should auto-run the {scenario_count} scenario(s) "
                "defined in project.yml. If Stage 4 finished without "
                "writing scenario results to the database, re-run Stage 4 "
                "or use the Scenario Runner page."
                if scenario_count
                else "No `scenarios:` block found in project.yml."
            ),
        )

    import geopandas as gpd

    gdf = bundle.results
    if not isinstance(gdf, gpd.GeoDataFrame):
        raise HTTPException(
            500,
            f"Artifact {bundle.results_artifact_id} is not a GeoDataFrame; "
            "scenario producer must include a geometry column.",
        )

    # Compute delta_Scenario{N} columns when missing.
    if PRED_BASELINE_COL in gdf.columns:
        for col in list(gdf.columns):
            if not col.startswith(PRED_SCENARIO_PREFIX):
                continue
            delta_col = col.replace(PRED_SCENARIO_PREFIX, DELTA_SCENARIO_PREFIX)
            if delta_col in gdf.columns:
                continue
            try:
                gdf[delta_col] = (
                    gdf[col].astype(float) - gdf[PRED_BASELINE_COL].astype(float)
                )
            except Exception:
                pass

    if gdf.crs is not None and str(gdf.crs) != "EPSG:4326":
        gdf = gdf.to_crs(epsg=4326)
    geojson_data = gdf.__geo_interface__

    summary_records: list[dict[str, Any]] = []
    if bundle.summary is not None:
        try:
            summary_records = bundle.summary.to_dict(orient="records")
        except Exception:
            summary_records = []

    return {
        "geojson": geojson_data,
        "summary": summary_records,
        "results_artifact_id": bundle.results_artifact_id,
        "summary_artifact_id": bundle.summary_artifact_id,
    }


@app.get("/results/scenarios/attribution")
async def get_scenario_attribution():
    """Return the per-scenario × per-variable attribution table."""
    from sparc.registry.store import ArtifactStore
    from sparc.scenario import SCENARIO_ATTRIBUTION, SCENARIO_STAGE

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id=SCENARIO_ATTRIBUTION, stage=SCENARIO_STAGE,
            hint="No active run registry.",
        )
    store = ArtifactStore(state.registry)
    if not store.has(SCENARIO_STAGE, SCENARIO_ATTRIBUTION):
        raise _missing_artifact_response(
            artifact_id=SCENARIO_ATTRIBUTION, stage=SCENARIO_STAGE,
            hint="Scenario simulator did not produce an attribution table.",
        )
    df = store.read_table(SCENARIO_STAGE, SCENARIO_ATTRIBUTION)
    return {"records": df.to_dict(orient="records"), "columns": list(df.columns)}


@app.get("/results/scenarios/trajectory")
async def get_scenario_trajectory(scenario_id: int | None = None):
    """Return the long-format scenario trajectory table, optionally filtered.

    Expected columns: ``scenario_id, t, geometry_id, value, std?``.
    """
    from sparc.registry.store import ArtifactStore
    from sparc.scenario import SCENARIO_STAGE, SCENARIO_TRAJECTORY

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id=SCENARIO_TRAJECTORY, stage=SCENARIO_STAGE,
            hint="No active run registry.",
        )
    store = ArtifactStore(state.registry)
    if not store.has(SCENARIO_STAGE, SCENARIO_TRAJECTORY):
        raise _missing_artifact_response(
            artifact_id=SCENARIO_TRAJECTORY, stage=SCENARIO_STAGE,
            hint="Scenario simulator did not produce a trajectory table.",
        )
    df = store.read_table(SCENARIO_STAGE, SCENARIO_TRAJECTORY)
    if scenario_id is not None and "scenario_id" in df.columns:
        df = df[df["scenario_id"] == scenario_id]
    return {"records": df.to_dict(orient="records"), "columns": list(df.columns)}


@app.get("/results/scenarios/uncertainty")
async def get_scenario_uncertainty(scenario_id: int | None = None):
    """Return per-feature uncertainty (mean/std/p05/p95) as GeoJSON.

    When ``scenario_id`` is provided, the per-scenario prediction column is
    included; otherwise only the uncertainty columns are exposed.
    """
    from sparc.registry.store import ArtifactStore
    from sparc.scenario import (
        UNCERTAINTY_COLS,
        ScenarioBundle,
        pred_column,
    )

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint="No active run registry.",
        )
    store = ArtifactStore(state.registry)
    bundle = ScenarioBundle.from_store(store)
    if bundle.results is None:
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint="Scenario results not in database.",
        )
    if not bundle.has_uncertainty():
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint=(
                "Active scenario_results table has no uncertainty columns "
                f"({', '.join(UNCERTAINTY_COLS)}). Re-run with the MC-Dropout "
                "ensemble enabled."
            ),
        )
    import geopandas as gpd

    gdf = bundle.results
    keep = [c for c in UNCERTAINTY_COLS if c in gdf.columns]
    if scenario_id is not None:
        sc = bundle.get_scenario(scenario_id)
        if sc is not None:
            keep.append(sc.pred_column)
            if sc.delta_column:
                keep.append(sc.delta_column)
    if isinstance(gdf, gpd.GeoDataFrame):
        out = gdf[[*keep, gdf.geometry.name]].copy()
        if out.crs is not None and str(out.crs) != "EPSG:4326":
            out = out.to_crs(epsg=4326)
        return {
            "geojson": out.__geo_interface__,
            "results_artifact_id": bundle.results_artifact_id,
        }
    raise HTTPException(500, "scenario results table is not a GeoDataFrame")


@app.get("/results/scenario_library/timeline")
async def get_scenario_library_timeline():
    """Return the scenario library as a flat parent/child timeline for the UI."""
    from sparc.registry.store import ArtifactStore
    from sparc.scenario import SCENARIO_LIBRARY, SCENARIO_STAGE

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id=SCENARIO_LIBRARY, stage=SCENARIO_STAGE,
            hint="No active run registry.",
        )
    store = ArtifactStore(state.registry)
    if not store.has(SCENARIO_STAGE, SCENARIO_LIBRARY):
        # Fall back to the on-disk library.jsonl (still the producer today)
        # but only as a read — no consumer writes there.
        try:
            from sparc.scenario.library import ScenarioTimeline
            tl = ScenarioTimeline.load(state.registry.output_dir)
            return {"entries": [e.model_dump() for e in tl.entries]}
        except Exception:
            raise _missing_artifact_response(
                artifact_id=SCENARIO_LIBRARY, stage=SCENARIO_STAGE,
                hint="No scenario library produced yet.",
            )
    payload = store.read_struct(SCENARIO_STAGE, SCENARIO_LIBRARY)
    return payload


@app.get("/results/report")
async def get_report_data():
    """Compile all stage results into a structured payload for the report view (DB-only)."""
    cfg = state.project_config or {}
    report: dict[str, Any] = {}

    # Project info
    raw = state.raw_project_yaml or {}
    report["project"] = raw.get("project", {})
    report["data_summary"] = state.data_summary or {}

    # Predictors
    predictors = cfg.get("predictors", {})
    if isinstance(predictors, list):
        report["predictors"] = predictors
    elif isinstance(predictors, dict):
        report["predictors"] = predictors.get("base_model", [])
    else:
        report["predictors"] = []

    # Causal + physics + pipeline config
    report["causal"] = cfg.get("causal", {})
    report["physics"] = cfg.get("physics", {})
    report["pipeline"] = cfg.get("pipeline", {})

    if state.registry is None:
        return report

    from sparc.registry.run_registry import set_active_registry
    from sparc.registry.store import ArtifactStore
    set_active_registry(state.registry)
    try:
        store = ArtifactStore(state.registry)

        # Stage 0 — correlogram summary
        if store.has("0", "correlogram_results"):
            corr_data = store.read_any("0", "correlogram_results")
            individual = (corr_data or {}).get("individual_results", {})
            report["correlogram"] = {
                var: {
                    "optimal_bandwidth": info.get("optimal_bandwidth"),
                    "effective_range": info.get("effective_range"),
                    "max_moran_i": info.get("max_moran_i"),
                }
                for var, info in individual.items()
            }

        # Stage 1 — GWEN variable importance
        if store.has("1", "gwen_variable_importance"):
            gwen_df = store.read_any("1", "gwen_variable_importance")
            if gwen_df is not None:
                report["gwen"] = gwen_df.to_dict(orient="records")

        # Stage 2 — spatial CV models
        if store.has("2", "oof_predictions"):
            oof_df = store.read_any("2", "oof_predictions")
            if oof_df is not None:
                report["spatial_cv_models"] = list(oof_df.columns)
        if store.has("2", "ensemble_results"):
            report["ensemble_results"] = store.read_any("2", "ensemble_results")

        # Stage 3 — causal coefficients + dose-response + propensity
        for art_id, key in (
            ("scenario_coefficients", "causal_results"),
            ("dose_response_curves", "dose_response"),
            ("propensity_diagnostics", "propensity_diagnostics"),
            ("causal_diagnostics", "causal_diagnostics"),
        ):
            if store.has("3", art_id):
                report[key] = store.read_any("3", art_id)

        # Stage 4 — scenario summary (mode-variant precedence)
        for variant in (
            "scenario_summary",
            "scenario_summary_dag",
            "scenario_summary_hybrid",
            "scenario_summary_reprediction",
            "scenario_mc_consensus_summary",
            "scenario_mc_consensus",
        ):
            if store.has("4", variant):
                summary_df = store.read_any("4", variant)
                if summary_df is not None and hasattr(summary_df, "to_dict"):
                    report["scenario_summary"] = summary_df.to_dict(orient="records")
                    break
    finally:
        set_active_registry(None)

    # Plot URLs intentionally omitted: visualisations are rendered live in
    # the desktop app from artifacts.db data and exported client-side.
    report["plots"] = {}
    return report


# ------------------------------------------------------------------
# Decision-support endpoints (Phase 9): rank candidate interventions by
# causal effect-per-cost with optional equity weights and uncertainty
# penalty.  Inputs come from existing scenario / NUTS results so the
# optimiser is fully causal (not correlation-based).
# ------------------------------------------------------------------

@app.get("/decision/candidates")
async def get_decision_candidates():
    """Build candidate interventions from scenario + NUTS results."""
    from sparc.decision import propose_candidates_from_scenarios

    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    scenarios: list[dict] = []
    nuts: dict | None = None
    try:
        detail = await get_scenario_detail()  # type: ignore[misc]
        if isinstance(detail, dict):
            scenarios = detail.get("summary") or []
    except HTTPException:
        scenarios = []
    try:
        nuts = await get_nuts_summary()  # type: ignore[misc]
    except HTTPException:
        nuts = None
    except NameError:
        nuts = None

    cands = propose_candidates_from_scenarios(scenarios, nuts)
    return {"candidates": [
        {
            "name": c.name, "treatment": c.treatment, "magnitude": c.magnitude,
            "mean_effect": c.mean_effect, "effect_std": c.effect_std,
            "cost": c.cost, "equity_weight": c.equity_weight, "notes": c.notes,
        }
        for c in cands
    ]}


@app.post("/decision/optimize")
async def post_decision_optimize(body: dict = Body(default_factory=dict)):
    """Rank intervention candidates and (optionally) pick a portfolio under a budget.

    Body fields (all optional)::

        {
          "candidates":          [InterventionCandidate, ...],   # explicit override
          "budget":              float | null,
          "robustness_lambda":   float (default 0),
          "minimise":            bool  (default false)
        }

    If ``candidates`` is omitted the endpoint derives them from the latest
    scenario + NUTS posterior outputs.
    """
    from sparc.decision import (
        InterventionCandidate, rank_interventions, propose_candidates_from_scenarios,
    )

    raw_candidates = body.get("candidates")
    if raw_candidates:
        candidates = [
            InterventionCandidate(
                name=str(r.get("name", "candidate")),
                treatment=str(r.get("treatment", r.get("name", "candidate"))),
                magnitude=float(r.get("magnitude", 0.0) or 0.0),
                mean_effect=float(r.get("mean_effect", r.get("delta", 0.0)) or 0.0),
                effect_std=float(r.get("effect_std", 0.0) or 0.0),
                cost=float(r.get("cost", 1.0) or 1.0),
                equity_weight=float(r.get("equity_weight", 1.0) or 1.0),
                notes=str(r.get("notes", "") or ""),
            )
            for r in raw_candidates
            if isinstance(r, dict)
        ]
    else:
        if state.project_config is None:
            raise HTTPException(400, "No project loaded and no candidates provided")
        scenarios: list[dict] = []
        nuts: dict | None = None
        try:
            detail = await get_scenario_detail()  # type: ignore[misc]
            if isinstance(detail, dict):
                scenarios = detail.get("summary") or []
        except HTTPException:
            scenarios = []
        try:
            nuts = await get_nuts_summary()  # type: ignore[misc]
        except (HTTPException, NameError):
            nuts = None
        candidates = propose_candidates_from_scenarios(scenarios, nuts)

    if not candidates:
        raise HTTPException(404, "No intervention candidates available")

    result = rank_interventions(
        candidates,
        budget=body.get("budget"),
        robustness_lambda=float(body.get("robustness_lambda", 0.0) or 0.0),
        minimise=bool(body.get("minimise", False)),
    )
    return {
        "ranked": result.ranked,
        "settings": result.settings,
        "equity_summary": result.equity_summary,
    }


@app.post("/decision/uncertainty")
async def post_decision_uncertainty(body: dict = Body(default_factory=dict)):
    """Monte-Carlo uncertainty for the decision optimizer.

    Body fields (same as ``/decision/optimize`` plus ``n_draws`` and ``seed``)::

        {
          "candidates":         [...],   # optional, falls back to scenarios+NUTS
          "budget":             float | null,
          "robustness_lambda":  float,
          "minimise":           bool,
          "n_draws":            int (default 500),
          "seed":               int | null (default 42)
        }

    Returns selection probabilities, rank quantiles, and effect quantiles
    per candidate.
    """
    from sparc.decision import (
        InterventionCandidate,
        propose_candidates_from_scenarios,
        monte_carlo_decision,
    )

    raw_candidates = body.get("candidates")
    if raw_candidates:
        candidates = [
            InterventionCandidate(
                name=str(r.get("name", "candidate")),
                treatment=str(r.get("treatment", r.get("name", "candidate"))),
                magnitude=float(r.get("magnitude", 0.0) or 0.0),
                mean_effect=float(r.get("mean_effect", r.get("delta", 0.0)) or 0.0),
                effect_std=float(r.get("effect_std", 0.0) or 0.0),
                cost=float(r.get("cost", 1.0) or 1.0),
                equity_weight=float(r.get("equity_weight", 1.0) or 1.0),
                notes=str(r.get("notes", "") or ""),
            )
            for r in raw_candidates
            if isinstance(r, dict)
        ]
    else:
        if state.project_config is None:
            raise HTTPException(400, "No project loaded and no candidates provided")
        scenarios: list[dict] = []
        nuts: dict | None = None
        try:
            detail = await get_scenario_detail()  # type: ignore[misc]
            if isinstance(detail, dict):
                scenarios = detail.get("summary") or []
        except HTTPException:
            scenarios = []
        try:
            nuts = await get_nuts_summary()  # type: ignore[misc]
        except (HTTPException, NameError):
            nuts = None
        candidates = propose_candidates_from_scenarios(scenarios, nuts)

    if not candidates:
        raise HTTPException(404, "No intervention candidates available")

    results = monte_carlo_decision(
        candidates,
        budget=body.get("budget"),
        robustness_lambda=float(body.get("robustness_lambda", 0.0) or 0.0),
        minimise=bool(body.get("minimise", False)),
        n_draws=int(body.get("n_draws", 500) or 500),
        seed=body.get("seed", 42),
    )
    return {
        "uncertainty": [
            {
                "candidate": r.candidate,
                "selection_probability": r.selection_probability,
                "rank_mean": r.rank_mean,
                "rank_p10": r.rank_p10,
                "rank_p90": r.rank_p90,
                "effect_p10": r.effect_p10,
                "effect_p50": r.effect_p50,
                "effect_p90": r.effect_p90,
            }
            for r in results
        ],
        "settings": {
            "n_draws": int(body.get("n_draws", 500) or 500),
            "robustness_lambda": float(body.get("robustness_lambda", 0.0) or 0.0),
            "budget": body.get("budget"),
            "minimise": bool(body.get("minimise", False)),
        },
    }


@app.get("/equity/census")
async def get_equity_census(
    minlon: float | None = Query(default=None),
    minlat: float | None = Query(default=None),
    maxlon: float | None = Query(default=None),
    maxlat: float | None = Query(default=None),
):
    """Auto-fetch area-wide ACS demographics for the project bbox.

    If bbox params are omitted, falls back to the loaded project's
    coordinate columns (``data.lon_column``/``lat_column``) by reading
    the source CSV. Only works for US locations (Census Geocoder).
    """
    from sparc.data.census_equity import (
        fetch_area_demographics,
        context_to_equity_layers,
    )

    # Fallback: derive bbox from the project's data file.
    if None in (minlon, minlat, maxlon, maxlat):
        if state.project_config is None:
            raise HTTPException(400, "No project loaded; provide bbox params or load a project.")
        try:
            data_cfg = state.project_config.get("data", {}) or {}
            csv_path = data_cfg.get("path") or data_cfg.get("source")
            lon_col = data_cfg.get("lon_column") or "lon"
            lat_col = data_cfg.get("lat_column") or "lat"
            if not csv_path or not Path(csv_path).exists():
                raise HTTPException(400, "Project data file not found; provide explicit bbox.")
            import pandas as pd
            df = pd.read_csv(csv_path, usecols=[lon_col, lat_col])
            minlon = float(df[lon_col].min())
            maxlon = float(df[lon_col].max())
            minlat = float(df[lat_col].min())
            maxlat = float(df[lat_col].max())
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(400, f"Could not derive bbox from project data: {exc}")

    ctx = fetch_area_demographics(
        minlon=float(minlon), minlat=float(minlat),
        maxlon=float(maxlon), maxlat=float(maxlat),
    )
    return {
        "bbox": {"minlon": minlon, "minlat": minlat, "maxlon": maxlon, "maxlat": maxlat},
        **context_to_equity_layers(ctx),
    }


@app.post("/decision/equity")
async def post_decision_equity(body: dict = Body(default_factory=dict)):
    """Combine equity layers into per-candidate weights.

    Body fields::

        {
          "candidate_names": [str, ...],
          "layers":          {layer_name: [float, ...]},   # aligned to names
          "weights":         {layer_name: float},          # optional convex weights
          "invert":          {layer_name: bool}            # flip 1−x for "advantage" layers
        }

    Returns ``{ "scores": [...], "disparity_index": float }``.
    """
    from sparc.decision import combine_equity_layers, disparity_index

    names = body.get("candidate_names") or []
    layers = body.get("layers") or {}
    if not names or not layers:
        raise HTTPException(400, "candidate_names and layers are required")

    try:
        scores = combine_equity_layers(
            layers,
            names,
            weights=body.get("weights"),
            invert=body.get("invert"),
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))

    return {
        "scores": [
            {
                "candidate": s.candidate,
                "weight": s.weight,
                "layer_breakdown": s.layer_breakdown,
            }
            for s in scores
        ],
        "disparity_index": disparity_index([s.weight for s in scores]),
    }


@app.get("/decision/targeting")
async def get_decision_targeting(
    variable: str = Query(...),
    top_k: int = Query(50, ge=1, le=2000),
):
    """Return per-location deployment priority as GeoJSON.

    Priority = |CATE| / cost  (cost defaults to 1).  The endpoint reuses
    the spatial CATE map for the requested treatment ``variable``,
    annotates each feature with a ``priority`` property, sorts the result
    descending, and flags the top ``top_k`` features with ``selected=True``.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    # Reuse the existing CATE map endpoint to load the GeoJSON.
    cate_geo = await get_cate_map(variable=variable)  # type: ignore[misc]
    if not isinstance(cate_geo, dict):
        raise HTTPException(500, "CATE map response was not a dict")
    features = cate_geo.get("features") or []
    if not features:
        raise HTTPException(404, f"No CATE features for '{variable}'")

    cate_field = f"cate_{variable}"
    enriched: list[tuple[float, dict]] = []
    for feat in features:
        props = dict(feat.get("properties") or {})
        try:
            cate_val = float(props.get(cate_field, 0.0) or 0.0)
        except (TypeError, ValueError):
            cate_val = 0.0
        # Optional cost field; defaults to 1 for uniform cost.
        try:
            cost = float(props.get("cost", 1.0) or 1.0)
        except (TypeError, ValueError):
            cost = 1.0
        priority = abs(cate_val) / max(cost, 1e-9)
        props["priority"] = priority
        props["abs_cate"] = abs(cate_val)
        new_feat = {**feat, "properties": props}
        enriched.append((priority, new_feat))

    enriched.sort(key=lambda pair: pair[0], reverse=True)
    cutoff = min(top_k, len(enriched))
    out_features: list[dict] = []
    priorities: list[float] = []
    for rank, (priority, feat) in enumerate(enriched, start=1):
        feat["properties"]["rank"] = rank
        feat["properties"]["selected"] = rank <= cutoff
        out_features.append(feat)
        priorities.append(priority)

    return {
        "type": "FeatureCollection",
        "features": out_features,
        "summary": {
            "variable": variable,
            "n_features": len(out_features),
            "top_k": cutoff,
            "max_priority": max(priorities) if priorities else 0.0,
            "min_priority": min(priorities) if priorities else 0.0,
        },
    }


# ------------------------------------------------------------------
# CATE map & increment endpoints (Phase 2)
# ------------------------------------------------------------------

def _read_cate_column_from_store(variable: str, column: str = "multiplier_mean"):
    """Return per-cell values of *column* for *variable* from the
    ``("3","cate_summary")`` long-format table, or ``None`` if absent.

    The v4 causal stack persists CATE results as a long-format table
    keyed by ``(cell_id, treatment)``. Filter by treatment, sort by
    ``cell_id``, return *column* as a NumPy array.

    The default ``multiplier_mean`` is the dimensionless spatial sensitivity
    multiplier consumed internally by the scenario simulator; pass
    ``column="cate_mean"`` to read the Bayesian per-cell coefficient β(s)
    in original ΔY-per-unit-T units (this is what the CATE map and the
    Decision Support benefit vector display).
    """
    if state.registry is None:
        return None
    try:
        from sparc.registry.run_registry import set_active_registry
        from sparc.registry.store import ArtifactStore
        import numpy as np

        set_active_registry(state.registry)
        try:
            store = ArtifactStore(state.registry)
            if not store.has("3", "cate_summary"):
                return None
            df = store.read_table("3", "cate_summary")
            if df is None or "treatment" not in df.columns:
                return None
            sub = df[df["treatment"] == variable]
            if sub.empty or column not in sub.columns:
                return None
            sub = sub.sort_values("cell_id")
            return np.asarray(sub[column].to_numpy(), dtype=float)
        finally:
            set_active_registry(None)
    except Exception:
        return None


def _read_cate_multiplier_from_store(variable: str):
    """Back-compat wrapper: return ``multiplier_mean`` (used by scenario simulator)."""
    return _read_cate_column_from_store(variable, "multiplier_mean")


def _read_cate_coefficient_from_store(variable: str, *, with_ci: bool = False):
    """Return the Bayesian per-cell coefficient β(s) for *variable*.

    Reads ``cate_mean`` from the ``("3","cate_summary")`` table. When
    ``with_ci=True``, also returns ``cate_ci5`` and ``cate_ci95`` so the
    caller can render uncertainty (e.g. stipple cells whose 90% CI crosses
    zero on the CATE map).

    Returns ``None`` if the artifact is missing, otherwise a dict with
    keys ``mean``, ``ci5``, ``ci95``, ``source``. ``ci5``/``ci95`` may be
    ``None`` when ``with_ci=False`` or when those columns are absent.
    """
    mean = _read_cate_column_from_store(variable, "cate_mean")
    if mean is None:
        return None
    src = _read_cate_column_from_store(variable, "source")
    source_label = None
    if src is not None and len(src):
        try:
            source_label = str(src[0])
        except Exception:
            source_label = None
    out: dict[str, Any] = {"mean": mean, "source": source_label}
    if with_ci:
        out["ci5"] = _read_cate_column_from_store(variable, "cate_ci5")
        out["ci95"] = _read_cate_column_from_store(variable, "cate_ci95")
    return out


@app.get("/results/causal/cate_map")
async def get_cate_map(
    variable: str = Query(...),
    with_uncertainty: bool = Query(
        False,
        description="When true, also emit per-cell coef_ci5_<var> / coef_ci95_<var>"
                    " properties so the frontend can stipple cells whose 90% CI"
                    " crosses zero (“not significant”).",
    ),
):
    """Return the Bayesian per-cell coefficient β(s) for *variable* as GeoJSON.

    DB-only: reads ``cate_mean`` (and optionally ``cate_ci5``/``cate_ci95``)
    from the ``("3","cate_summary")`` table and joins onto the in-memory
    project geometry. β(s) is the posterior mean local treatment effect
    in original ΔY-per-unit-T units, produced by ``BayesianSpatialCATE``
    (NUTS over a random-Fourier-feature basis on coordinates).

    GeoJSON properties are written under both ``coef_<var>`` (canonical) and
    ``cate_<var>`` (one-release alias for back-compat with the old multiplier
    payload). Top-level foreign members ``units``, ``source``, and ``variable``
    are added so the desktop can render the legend and equation panel.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    payload = _read_cate_coefficient_from_store(variable, with_ci=with_uncertainty)
    if payload is None:
        raise _missing_artifact_response(
            artifact_id=f"cate_summary[{variable}]", stage="3",
            hint=(
                f"No CATE map for variable '{variable}'. Re-run Stage 3 with "
                "`causal.estimate_cate: true` and at least one treatment in your DAG."
            ),
        )

    if state.data is None or not hasattr(state.data, "geometry"):
        raise HTTPException(
            500,
            "Project data has no geometry; cannot project CATE coefficients "
            "onto the map. Re-load the project so the input dataset is"
            " parsed with coordinate columns.",
        )

    import geopandas as gpd

    src = state.data
    coef = payload["mean"]
    n = min(len(src), len(coef))
    cols: dict[str, Any] = {
        f"coef_{variable}": coef[:n],
        f"cate_{variable}": coef[:n],  # back-compat alias — same values as coef_<var>
    }
    if with_uncertainty:
        ci5 = payload.get("ci5")
        ci95 = payload.get("ci95")
        if ci5 is not None and ci95 is not None:
            cols[f"coef_ci5_{variable}"] = ci5[:n]
            cols[f"coef_ci95_{variable}"] = ci95[:n]
            # “Not significant at 90%” flag — the frontend uses this to
            # dim/stipple cells whose CI brackets zero.
            import numpy as _np
            cols[f"coef_ns_{variable}"] = (
                (_np.asarray(ci5[:n], dtype=float) <= 0.0)
                & (_np.asarray(ci95[:n], dtype=float) >= 0.0)
            )

    gdf = gpd.GeoDataFrame(cols, geometry=src.geometry.values[:n], crs=src.crs)
    if gdf.crs is not None and str(gdf.crs) != "EPSG:4326":
        gdf = gdf.to_crs(epsg=4326)
    fc = gdf.__geo_interface__
    # GeoJSON foreign members — RFC 7946 §6.1.
    fc["variable"] = variable
    fc["units"] = "ΔY per unit T (posterior mean of β(s))"
    fc["source"] = payload.get("source") or "bayesian_nuts"
    fc["with_uncertainty"] = bool(with_uncertainty)
    return fc


@app.get("/results/causal/cate_map/variables")
async def get_cate_variables():
    """Return list of variables that have spatial CATE multiplier maps."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    from sparc.run.pipeline_paths import PipelinePaths

    try:
        paths = PipelinePaths.from_config(state.project_config)
    except Exception:
        raise HTTPException(404, "Cannot resolve output paths")

    variables: list[str] = []
    diagnostics: dict[str, Any] = {}

    # 1. Prefer ("3","cate_summary") long-format table.
    if state.registry is not None:
        try:
            from sparc.registry.run_registry import set_active_registry
            from sparc.registry.store import ArtifactStore
            set_active_registry(state.registry)
            try:
                store = ArtifactStore(state.registry)
                if store.has("3", "cate_summary"):
                    df = store.read_table("3", "cate_summary")
                    if df is not None and "treatment" in df.columns:
                        variables.extend(
                            sorted(str(t) for t in df["treatment"].unique())
                        )
                        diagnostics["cate_summary_hits"] = len(variables)
            finally:
                set_active_registry(None)
        except Exception as exc:
            diagnostics["cate_summary_error"] = str(exc)

    # 2. Legacy registry: any cate_multiplier::* artifact wins.
    if not variables and state.registry is not None:
        for art in state.registry.list_for_stage("3"):
            if art.id.startswith("cate_multiplier::") and not art.partial:
                var = art.metadata.get("variable") or art.id.split("::", 1)[1]
                variables.append(var)
        diagnostics["registry_hits"] = len(variables)

    # 3. Fallback: per-variable .npy files on disk.
    # (DB-only: legacy disk lookups removed in v4 refresh.)

    # 3. Fallback: column names in spatial_cate_maps.gpkg.
    # (DB-only: legacy gpkg lookups removed in v4 refresh.)

    # 4. Empty-state hint to help the user act on the Budget Optimizer screen.
    sorted_vars = sorted(set(variables))
    payload: dict[str, Any] = {"variables": sorted_vars}
    if not sorted_vars:
        causal_cfg = state.project_config.get("causal", {}) or {}
        inference = (causal_cfg.get("inference") or "").lower()
        estimate_cate = causal_cfg.get("estimate_cate", True)
        reasons: list[str] = []
        if not estimate_cate:
            reasons.append("`causal.estimate_cate` is false in project.yml.")
        if inference == "bayesian":
            reasons.append(
                "Bayesian inference path historically skipped CATE; this is "
                "fixed in newer Stage 3 runs but the existing run was "
                "produced before the fix."
            )
        if not reasons:
            reasons.append(
                "Stage 3 finished but did not write spatial CATE multipliers. "
                "Check the DAG for at least one treatment node."
            )
        payload["empty_reason"] = " ".join(reasons)
        payload["next_action"] = (
            "Re-run Stage 3 with `causal.estimate_cate: true` and at least "
            "one treatment in your DAG."
        )
    if diagnostics:
        payload["diagnostics"] = diagnostics
    return payload


# NOTE: the legacy ``/results/local_coefficients`` and
# ``/results/local_coefficients/variables`` endpoints have been removed.
# MGWR/GWR are kept as Stage-2 diagnostics, but their coefficients are
# no longer surfaced for effect or decision logic — every consumer now
# reads the Bayesian per-cell coefficient β(s) from
# ``("3","cate_summary")`` via :func:`_read_cate_coefficient_from_store`.


@app.get("/results/scenarios/variables")
async def get_scenario_variables():
    """Return the list of scenario variables and their available increments.

    Response shape::

        {
            "variables": {
                "Pct_Canopy": { "increments": [5, 10, 15, ...], "sign": "plus" },
                "Pct_Impervious": { "increments": [-5, -10, ...], "sign": "minus" }
            }
        }

    DB-only: reads ``scenario_summary*`` from ``artifacts.db`` via the
    ``ArtifactStore``. The on-disk ``scenario_summary*.csv`` files are
    deliberately ignored.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(404, "No active run registry")

    from sparc.registry.store import ArtifactStore
    from sparc.scenario import SCENARIO_STAGE, SCENARIO_SUMMARY_VARIANTS

    store = ArtifactStore(state.registry)
    summary_df = None
    for variant in SCENARIO_SUMMARY_VARIANTS:
        if store.has(SCENARIO_STAGE, variant):
            summary_df = store.read_table(SCENARIO_STAGE, variant)
            break
    if summary_df is None:
        raise HTTPException(404, "No scenario summary found in artifacts.db")

    variables: dict[str, dict] = {}
    for _, row in summary_df.iterrows():
        var = str(row.get("Variable", ""))
        inc = row.get("Increment")
        if not var or inc is None:
            continue
        try:
            inc = float(inc)
        except (TypeError, ValueError):
            continue
        if var not in variables:
            variables[var] = {"increments": [], "sign": "plus" if inc >= 0 else "minus"}
        if inc not in variables[var]["increments"]:
            variables[var]["increments"].append(inc)

    for info in variables.values():
        info["increments"].sort(key=lambda x: abs(x))

    return {"variables": variables}


@app.get("/results/scenarios/increment")
async def get_scenario_increment(variable: str = Query(...), increment: float = Query(...)):
    """Return GeoJSON filtered to a specific variable+increment scenario.

    DB-only: reads ``scenario_summary*`` and ``scenario_results*`` from
    ``artifacts.db`` via the ``ArtifactStore``. Legacy ``.csv`` / ``.gpkg``
    files in ``Stage_4_Scenarios`` are intentionally not consulted.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(404, "No active run registry")

    from sparc.registry.store import ArtifactStore
    from sparc.scenario import (
        SCENARIO_STAGE,
        SCENARIO_SUMMARY_VARIANTS,
        SCENARIO_RESULTS_VARIANTS,
    )

    store = ArtifactStore(state.registry)

    summary_df = None
    for variant in SCENARIO_SUMMARY_VARIANTS:
        if store.has(SCENARIO_STAGE, variant):
            summary_df = store.read_table(SCENARIO_STAGE, variant)
            break
    if summary_df is None:
        raise HTTPException(404, "No scenario summary found in artifacts.db")

    mask = (
        (summary_df["Variable"].str.lower() == variable.lower()) &
        (summary_df["Increment"].astype(float).round(6) == round(float(increment), 6))
    )
    matched = summary_df[mask]
    if matched.empty:
        raise HTTPException(404, f"No scenario found for {variable} at increment {increment}")

    gdf = None
    for variant in SCENARIO_RESULTS_VARIANTS:
        if store.has(SCENARIO_STAGE, variant):
            gdf = store.read_table(SCENARIO_STAGE, variant)
            break
    if gdf is None:
        raise HTTPException(404, "No spatial scenario results found in artifacts.db")

    # ArtifactStore returns either a GeoDataFrame or a plain DataFrame — for
    # tables registered with ``geometry_col`` it's a GeoDataFrame already.
    import geopandas as gpd  # local import to keep top-level import light
    if not isinstance(gdf, gpd.GeoDataFrame) and "geometry" in gdf.columns:
        gdf = gpd.GeoDataFrame(gdf, geometry="geometry", crs="EPSG:4326")

    scenario_label = matched.iloc[0].get("Scenario", "")
    pred_col = f"pred_{scenario_label}"
    delta_col = f"delta_{scenario_label}"
    baseline_col = "pred_baseline" if "pred_baseline" in gdf.columns else None

    cols_to_keep = ["geometry"]
    if baseline_col:
        cols_to_keep.append(baseline_col)
    if pred_col in gdf.columns:
        cols_to_keep.append(pred_col)
    if delta_col in gdf.columns:
        cols_to_keep.append(delta_col)
    elif pred_col in gdf.columns and baseline_col:
        gdf[delta_col] = gdf[pred_col].astype(float) - gdf[baseline_col].astype(float)
        cols_to_keep.append(delta_col)

    result_gdf = gdf[cols_to_keep].copy()
    if isinstance(result_gdf, gpd.GeoDataFrame):
        if result_gdf.crs is not None and str(result_gdf.crs) != "EPSG:4326":
            result_gdf = result_gdf.to_crs(epsg=4326)

    return {
        "geojson": result_gdf.__geo_interface__,
        "summary": matched.to_dict(orient="records"),
    }


# ------------------------------------------------------------------
# Run registry (artifacts manifest) endpoints
# These MUST be declared before the parameterized /results/{stage}
# route below — FastAPI matches routes in declaration order, and the
# parameterized handler types `stage: int`, which would otherwise
# swallow /results/manifest with a 422.
# ------------------------------------------------------------------

@app.get("/results/manifest")
async def get_results_manifest(
    refresh: bool = Query(False, description="Reload from disk before returning."),
    rescan: bool = Query(False, description="Re-run the migrate_from_disk scan."),
):
    """Return the full run manifest (Pydantic ``RunManifest``).

    Frontend uses this to decide which Results panels to render and to
    surface 'unavailable, produced by stage X' empty-states.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        _attach_registry(state.project_config)
    if state.registry is None:
        raise HTTPException(503, "Run registry unavailable")

    if refresh or rescan:
        state.registry.load()
        if rescan:
            from sparc.run.pipeline_paths import PipelinePaths
            paths = PipelinePaths.from_config(state.project_config)
            state.registry.migrate_from_disk(paths)

    return state.registry.manifest.model_dump(mode="json", exclude_none=False)


@app.get("/results/manifest/{stage}")
async def get_stage_manifest(stage: str):
    """Return the per-stage manifest fragment (artifacts + status)."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        _attach_registry(state.project_config)
    if state.registry is None:
        raise HTTPException(503, "Run registry unavailable")

    sm = state.registry.manifest.stages.get(str(stage))
    if sm is None:
        return {"stage": str(stage), "status": "pending", "artifacts": {}}
    return sm.model_dump(mode="json", exclude_none=False)


@app.post("/results/manifest/rescan")
async def rescan_manifest():
    """Re-walk the run directory and import any unregistered artifacts."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        _attach_registry(state.project_config)
    if state.registry is None:
        raise HTTPException(503, "Run registry unavailable")
    from sparc.run.pipeline_paths import PipelinePaths
    paths = PipelinePaths.from_config(state.project_config)
    n = state.registry.migrate_from_disk(paths)
    return {"newly_registered": n,
            "total_artifacts": len(state.registry.manifest.all_artifacts())}


# ------------------------------------------------------------------
# Artifact fetch endpoints (db-resident artifacts via ArtifactStore)
#
# These render artifacts on demand from artifacts.db. The pipeline stops
# producing CSV / PNG / JSON files; users download them through here.
# ------------------------------------------------------------------

_RENDER_MIME = {
    "csv":     "text/csv",
    "json":    "application/json",
    "geojson": "application/geo+json",
    "html":    "text/html",
    "pkl":     "application/octet-stream",
    "joblib":  "application/octet-stream",
    "pt":      "application/octet-stream",
    "bin":     "application/octet-stream",
}
# NOTE: PNG / image MIME types are intentionally absent. The server
# never renders or serves rasterised images. The desktop app renders
# visualisations live from artifacts.db and exports them client-side.


def _ensure_registry_attached() -> None:
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        _attach_registry(state.project_config)
    if state.registry is None:
        raise HTTPException(503, "Run registry unavailable")


@app.get("/artifacts/{stage}/{artifact_id}.csv")
async def get_artifact_csv(stage: str, artifact_id: str, index: bool = False):
    """Render a registered artifact as CSV bytes."""
    _ensure_registry_attached()
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.report.render import render_csv, RenderError

    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            data = render_csv(stage, artifact_id, index=index)
        except RenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)

    return Response(
        content=data,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{artifact_id}.csv"'},
    )


@app.get("/artifacts/{stage}/{artifact_id}.json")
async def get_artifact_json(stage: str, artifact_id: str):
    """Render a registered artifact (struct or table) as JSON."""
    _ensure_registry_attached()
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.report.render import render_json, RenderError

    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            data = render_json(stage, artifact_id)
        except RenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)
    return Response(content=data, media_type="application/json")


@app.get("/artifacts/{stage}/{artifact_id}.geojson")
async def get_artifact_geojson(stage: str, artifact_id: str):
    """Render a geometry-bearing table as GeoJSON."""
    _ensure_registry_attached()
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.report.render import render_geojson, RenderError

    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            data = render_geojson(stage, artifact_id)
        except RenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)
    return Response(content=data, media_type="application/geo+json")


@app.get("/artifacts/{stage}/{artifact_id}.png")
async def get_artifact_png(stage: str, artifact_id: str, dpi: int = 150):
    """Render a registered artifact as a PNG via the figures module.

    Dispatches through ``sparc.report.figures.render_for_artifact``; returns
    404 when no renderer is registered for ``(stage, artifact_id)``.
    """
    _ensure_registry_attached()
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    try:
        from sparc.report.figures import FigureRenderError, render_for_artifact
    except ImportError as exc:
        raise HTTPException(503, f"figures module unavailable: {exc}")

    if state.registry is None:
        raise _missing_artifact_response(
            artifact_id=artifact_id, stage=stage,
            hint="No active run registry.",
        )
    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            data = render_for_artifact(stage, artifact_id, registry=state.registry, dpi=dpi)
        except FigureRenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)
    return Response(content=data, media_type="image/png")


# Catch-all native route — declared LAST so the suffixed routes above
# match first. Otherwise ``{artifact_id}`` swallows ``foo.csv`` etc. and
# the suffixed routes become dead code.
@app.get("/artifacts/{stage}/{artifact_id}")
async def get_artifact_native(stage: str, artifact_id: str):
    """Return an artifact in its native format (CSV for tables, JSON for structs, raw bytes for blobs)."""
    _ensure_registry_attached()
    # Bind the active registry so render_* helpers can find it.
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.report.render import render_native, RenderError

    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            data, ext = render_native(stage, artifact_id)
        except RenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)

    return Response(
        content=data,
        media_type=_RENDER_MIME.get(ext, "application/octet-stream"),
        headers={
            "Content-Disposition": f'attachment; filename="{artifact_id}.{ext}"',
        },
    )


@app.get("/results/bundle")
async def get_results_bundle():
    """Stream a ZIP of every registered artifact (data formats + manifest).

    Each artifact is rendered in its native format (CSV/JSON/GeoJSON) plus a
    PNG when a renderer is registered. The full ``RunManifest`` is included
    as ``manifest.json`` for reproducibility.
    """
    _ensure_registry_attached()
    if state.registry is None:
        raise HTTPException(404, "No active run registry.")

    import io
    import zipfile
    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.registry.store import ArtifactStore
    from sparc.report.render import (
        RenderError,
        render_csv,
        render_geojson,
        render_json,
    )

    try:
        from sparc.report.figures import FigureRenderError, render_for_artifact
        figures_available = True
    except ImportError:
        figures_available = False

    store = ArtifactStore(state.registry)
    manifest = state.registry.manifest

    buf = io.BytesIO()
    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(
                "manifest.json",
                manifest.model_dump_json(indent=2),
            )
            # Snapshot stage / artifact lists before iterating: figure
            # rendering may write a cached PNG back into the registry, which
            # mutates ``manifest.stages`` mid-loop and raises
            # ``RuntimeError: dictionary changed size during iteration``.
            stage_items = list(manifest.stages.items())
            for stage_id, stage_manifest in stage_items:
                artifact_items = list(stage_manifest.artifacts.items())
                for artifact_id, entry in artifact_items:
                    base = f"stage_{stage_id}/{artifact_id}"
                    # Native data export.
                    try:
                        if entry.storage_kind == "table":
                            geom = (entry.metadata or {}).get("geometry_col")
                            if geom:
                                zf.writestr(
                                    f"{base}.geojson",
                                    render_geojson(stage_id, artifact_id),
                                )
                            else:
                                zf.writestr(
                                    f"{base}.csv",
                                    render_csv(stage_id, artifact_id),
                                )
                        elif entry.storage_kind == "struct":
                            zf.writestr(
                                f"{base}.json",
                                render_json(stage_id, artifact_id),
                            )
                        # Blobs are skipped — they may be huge and opaque.
                    except RenderError as exc:
                        zf.writestr(
                            f"{base}.SKIPPED.txt",
                            f"render failed: {exc}".encode("utf-8"),
                        )
                    # Optional PNG.
                    if figures_available:
                        try:
                            png = render_for_artifact(stage_id, artifact_id, registry=state.registry)
                            zf.writestr(f"{base}.png", png)
                        except FigureRenderError:
                            pass
                        except Exception as exc:  # noqa: BLE001
                            zf.writestr(
                                f"{base}.png.SKIPPED.txt",
                                f"png render failed: {exc}".encode("utf-8"),
                            )
    finally:
        set_active_registry(prev)

    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": 'attachment; filename="sparc_results_bundle.zip"',
        },
    )


# ------------------------------------------------------------------
# Results endpoints (parameterized — MUST come after named routes above)
# ------------------------------------------------------------------

@app.get("/results/{stage:int}")
async def get_results(stage: int, format: str = Query("json", regex="^(json|geojson)$")):
    """Return in-memory results for a completed pipeline stage (DB-only).

    For artifact-style data use the dedicated `/results/<name>` endpoints.
    This endpoint only exposes the live in-memory result captured during
    the current process run.
    """
    result = state.get_result(stage)
    if result is None:
        raise _missing_artifact_response(
            artifact_id=f"stage_{stage}_result", stage=str(stage),
            hint=(
                f"No in-memory result for stage {stage}. Run the stage in this "
                "session, or query a specific artifact via `/results/<name>`."
            ),
        )
    if format == "geojson":
        return _to_geojson(result)
    return _to_json(result)


@app.get("/results/{stage:int}/predictions")
async def get_predictions(stage: int, format: str = Query("geojson", regex="^(json|geojson)$")):
    """Return spatial predictions for a stage (DB-only).

    Mapping:
      - stage 2 -> artifact ("2", "spatial_cv_predictions")
      - stage 4 -> ScenarioBundle baseline geometry from ("4", "scenario_results")
      - other stages do not produce per-point predictions.
    """
    if stage == 2:
        store = _open_store()
        try:
            if not store.has("2", "spatial_cv_predictions"):
                raise _missing_artifact_response(
                    artifact_id="spatial_cv_predictions", stage="2",
                    hint="Stage 2 has not produced spatial_cv_predictions. Run Stage 2.",
                )
            gdf = store.read_any("2", "spatial_cv_predictions")
        finally:
            from sparc.registry.run_registry import set_active_registry
            set_active_registry(None)
        if format == "geojson":
            return _to_geojson(gdf)
        return _to_json(gdf)

    if stage == 4:
        # Scenario predictions live under /results/scenarios/detail (full
        # mode-variant precedence + per-scenario column resolution).
        raise _missing_artifact_response(
            artifact_id="scenario_results", stage="4",
            hint="Use /results/scenarios/detail for stage 4 scenario predictions.",
        )

    hints = {
        0: "Stage 0 (EDA) does not write per-point predictions.",
        1: "Stage 1 (GWEN) writes feature-selection results, not predictions.",
        3: "Stage 3 produces causal effects, not direct predictions. View predictions on the Stage 2 tab.",
    }
    raise _missing_artifact_response(
        artifact_id=f"stage_{stage}_predictions", stage=str(stage),
        hint=hints.get(stage, f"No predictions registered for stage {stage}."),
    )


# Plot-image endpoints removed (Phase E): the server no longer serves
# rasterised images. Visualisations are rendered live in the desktop
# app from artifacts.db data and exported client-side.


# ------------------------------------------------------------------
# Scenario endpoints
# ------------------------------------------------------------------

@app.post("/scenarios/run")
async def run_scenarios():
    """Execute all scenarios defined in the loaded project."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    scenarios = state.project_config.get("scenarios", [])
    interaction_scenarios = state.project_config.get("interaction_scenarios", [])
    if not scenarios and not interaction_scenarios:
        raise HTTPException(400, "No scenarios defined in project.yml")

    # Run synchronously for now; a production version would use
    # background tasks or the WebSocket streaming endpoint.
    from sparc.interventions.scenario_simulator import ScenarioSimulator
    import pandas as pd

    config = state.project_config
    sim = ScenarioSimulator(config)
    sim.load_models()

    csv_path = config["paths"]["raw_csv_path"]
    data = pd.read_csv(csv_path)

    dag_file = config.get("causal", {}).get("dag_file")
    has_dag = dag_file and Path(dag_file).exists()
    requested_mode = config.get("pipeline", {}).get("scenario_mode", "auto")

    # Translate legacy mode aliases (one-shot deprecation warning).
    _LEGACY_MODE_ALIASES = {
        "physics":            "mode_1_physics",
        "dag_coefficient":    "mode_2_dag_local",
        "model_reprediction": "mode_3_full_ensemble",
        "hybrid":             "mode_4_hybrid",
    }
    scenario_mode = requested_mode
    if requested_mode == "bayesian":
        raise RuntimeError(
            "scenario_mode='bayesian' was removed in SPARC v4. "
            "Use 'mode_3_full_ensemble' or 'auto'."
        )
    if requested_mode in _LEGACY_MODE_ALIASES:
        new_key = _LEGACY_MODE_ALIASES[requested_mode]
        import warnings as _warnings
        _warnings.warn(
            f"scenario_mode='{requested_mode}' is deprecated; use '{new_key}'.",
            DeprecationWarning, stacklevel=2,
        )
        scenario_mode = new_key

    if scenario_mode == "auto":
        from sparc.__main__ import _resolve_auto_scenario_mode
        scenario_mode = _resolve_auto_scenario_mode(has_dag=has_dag)

    if scenario_mode == "mode_4_hybrid":
        summary_df, results_gdf = sim.run_with_hybrid_reprediction(data, verbose=True)
    elif scenario_mode == "mode_3_full_ensemble":
        summary_df, results_gdf = sim.run_with_model_reprediction(data, verbose=True)
    elif scenario_mode == "mode_2_dag_local":
        if has_dag:
            summary_df, results_gdf = sim.run_with_causal_dag(data, verbose=True)
        else:
            scenario_mode = "mode_1_physics"
            summary_df, results_gdf = sim.run(verbose=True)
    elif scenario_mode == "mode_1_physics":
        summary_df, results_gdf = sim.run(verbose=True)
    else:
        raise ValueError(
            f"Unknown scenario_mode '{scenario_mode}'. "
            f"Valid: auto, mode_1_physics, mode_2_dag_local, "
            f"mode_3_full_ensemble, mode_4_hybrid."
        )

    # --- Conservation checks on scenario results ---------------------
    conservation_violations = []
    try:
        import numpy as np
        from sparc.interventions.physics_priors import ConservationChecker
        checker = ConservationChecker()
        for scenario in scenarios:
            var = scenario['variable']
            for inc in scenario.get('increments', []):
                direction = scenario.get('direction', 'increase')
                delta_signed = -inc if direction == 'decrease' else inc
                col_label = f"total_{var}_{'minus' if delta_signed < 0 else 'plus'}_{str(inc).replace('.', 'p')}"
                if hasattr(results_gdf, 'columns') and col_label in results_gdf.columns:
                    deltas = {var: np.full(len(data), delta_signed)}
                    target_deltas = results_gdf[col_label].values
                    violations = checker.check(data, deltas, target_deltas=target_deltas, verbose=True)
                    conservation_violations.extend(violations)
    except Exception as e:
        print(f"  [CONSERVATION] Check skipped ({e})")

    state.store_result(4, {"summary": summary_df, "spatial": results_gdf})

    return {
        "status": "complete",
        "n_scenarios": len(scenarios) + len(interaction_scenarios),
        "summary_rows": len(summary_df),
        "scenario_mode": scenario_mode,
        "conservation_violations": len(conservation_violations),
    }


@app.get("/scenarios/results")
async def scenario_results(format: str = Query("geojson", regex="^(json|geojson)$")):
    """Return scenario simulation results."""
    result = state.get_result(4)
    if result is None:
        raise HTTPException(404, "No scenario results. Run /scenarios/run first.")

    if format == "geojson" and "spatial" in result:
        return _to_geojson(result["spatial"])

    if "summary" in result:
        return _to_json(result["summary"])

    return _to_json(result)


# ------------------------------------------------------------------
# Budget-constrained intervention allocation
# ------------------------------------------------------------------

class BudgetOptimizeRequest(BaseModel):
    """User-facing request for the budget optimizer.

    Source semantics
    ----------------
    `benefit_source` decides where per-cell unit benefits come from:

    - "cate"           — Bayesian per-cell coefficient β(s) for `variable`
                         (posterior mean of NUTS spatial CATE). This is the
                         only causal effect surface available; correlation-based
                         MGWR/GWR coefficients are no longer used here.
    - "uniform"        — every cell has unit benefit 1.0. Useful only as
                         a sanity check / equity reference.
    """

    budget: float = Field(..., gt=0, description="Total budget cap in cost units.")
    benefit_source: Literal["cate", "uniform"] = "cate"
    variable: Optional[str] = Field(
        None,
        description="Treatment variable (required for cate source).",
    )
    cost_column: Optional[str] = Field(
        None,
        description="Column in the active dataset holding per-cell unit cost.",
    )
    x_max_column: Optional[str] = Field(
        None,
        description="Column with per-cell maximum allocation (defaults to 1.0).",
    )
    solver: Literal["greedy", "greedy_2opt", "milp", "auto"] = "auto"
    pareto_sweep: bool = Field(
        True,
        description="Whether to additionally sweep the budget for the Pareto frontier.",
    )


def _resolve_benefits(
    source: str,
    variable: Optional[str],
) -> tuple[np.ndarray, str]:
    """Pull a per-cell benefit vector from the configured source (DB-only)."""

    if source == "uniform":
        # Uniform requires a known cell count — derive from the in-memory
        # project dataset. No disk reads.
        if state.data is None:
            raise HTTPException(
                400,
                "Cannot determine cell count: project data not loaded.",
            )
        n = len(state.data)
        return np.ones(n, dtype=float), "uniform unit benefits (1.0 per cell)"

    if not variable:
        raise HTTPException(400, f"benefit_source='{source}' requires `variable`.")

    if state.project_config is None:
        raise HTTPException(400, "No project loaded.")

    if source == "cate":
        # Bayesian per-cell coefficient β(s) from ("3","cate_summary").cate_mean.
        vals = _read_cate_column_from_store(variable, "cate_mean")
        if vals is None:
            raise HTTPException(
                404,
                f"Bayesian CATE for '{variable}' not available — run Stage 3 "
                "with `causal.estimate_cate: true` and at least one treatment.",
            )
        return vals, f"Bayesian local coefficient β(s) for '{variable}'"

    raise HTTPException(400, f"Unknown benefit_source: {source}")


def _resolve_per_cell_column(column: Optional[str], n_cells: int, default: float) -> np.ndarray:
    """Read a per-cell vector from the active dataset's CSV, or fall back to default."""
    if column is None:
        return np.full(n_cells, default, dtype=float)
    if state.project_config is None:
        return np.full(n_cells, default, dtype=float)
    cfg = state.project_config
    data_file = cfg.get("data", {}).get("file_path", cfg.get("paths", {}).get("raw_csv_path"))
    if not data_file:
        return np.full(n_cells, default, dtype=float)
    import pandas as pd
    df = pd.read_csv(data_file)
    if column not in df.columns:
        raise HTTPException(404, f"Column '{column}' not in dataset.")
    vals = df[column].to_numpy(dtype=float)[:n_cells]
    if len(vals) < n_cells:
        # Pad with default for any missing rows
        vals = np.concatenate([vals, np.full(n_cells - len(vals), default)])
    return vals


@app.post("/scenarios/budget/optimize")
async def budget_optimize(req: BudgetOptimizeRequest):
    """Allocate `req.budget` across cells to maximise expected benefit.

    Returns the allocation, summary statistics (total benefit/cost, Gini,
    counts of treated/fully-treated cells), the solver used, and — if
    requested — a Pareto sweep at multiple budget multipliers.

    The benefit signal is taken from the user-selected source
    (defensibility order: CATE > uniform). Cost defaults to 1.0 per cell;
    pass `cost_column` to use a real cost surface.
    """
    from sparc.scenario.budget import optimize as _opt, pareto_sweep as _sweep

    benefits, benefit_desc = _resolve_benefits(req.benefit_source, req.variable)
    n_cells = int(benefits.size)
    if n_cells == 0:
        raise HTTPException(404, "No cells in benefit source.")

    costs = _resolve_per_cell_column(req.cost_column, n_cells, default=1.0)
    x_max = _resolve_per_cell_column(req.x_max_column, n_cells, default=1.0)

    result = _opt(
        benefits=benefits,
        budget=req.budget,
        costs=costs,
        x_max=x_max,
        solver=req.solver,
    )

    out: dict = {
        "result": result.to_dict(),
        "n_cells": n_cells,
        "benefit_source": req.benefit_source,
        "benefit_description": benefit_desc,
    }

    if req.pareto_sweep:
        sweep = _sweep(
            benefits=benefits,
            budget=req.budget,
            costs=costs,
            x_max=x_max,
            solver=req.solver,
        )
        out["pareto"] = sweep.to_dict()

    return out


# ------------------------------------------------------------------
# Physics defaults endpoint
# ------------------------------------------------------------------

@app.get("/physics/defaults")
async def physics_defaults():
    """Return default literature-based physics priors."""
    try:
        from sparc.interventions.physics_priors import PhysicsPriors
        pp = PhysicsPriors()
        return {
            k: c.to_dict() for k, c in pp.coefficients.items()
        }
    except Exception as e:
        raise HTTPException(500, str(e))


# ------------------------------------------------------------------
# Data Processing endpoints
# ------------------------------------------------------------------

@app.post("/data/fishnet")
async def create_fishnet_endpoint(body: dict = Body(...)):
    """Create a fishnet grid from bounds + resolution, optionally clip to boundary."""
    from sparc.data.processing import create_fishnet, clip_to_boundary

    bounds = body.get("bounds")  # [minx, miny, maxx, maxy]
    resolution = body.get("resolution", 100)
    crs = body.get("crs", "EPSG:4326")

    if not bounds or len(bounds) != 4:
        raise HTTPException(400, "bounds must be [minx, miny, maxx, maxy]")

    gdf = create_fishnet(tuple(bounds), resolution, crs)

    boundary_path = body.get("boundary_path")
    if boundary_path and Path(boundary_path).exists():
        import geopandas as _gpd
        boundary = _gpd.read_file(boundary_path)
        if boundary.crs and boundary.crs.to_string() != crs:
            boundary = boundary.to_crs(crs)
        gdf = clip_to_boundary(gdf, boundary)

    # Persist to project directory
    if state.project_config:
        out_dir = Path(state.project_config.get("paths", {}).get("project_dir", "."))
        out_path = out_dir / "fishnet.gpkg"
        gdf.to_file(out_path, driver="GPKG")

    return {"n_cells": len(gdf), "columns": list(gdf.columns)}


@app.post("/data/zonal_stats")
async def zonal_stats_endpoint(body: dict = Body(...)):
    """Compute zonal statistics for raster layers on a fishnet."""
    from sparc.data.processing import run_zonal_stats
    import geopandas as _gpd

    fishnet_path = body.get("fishnet_path")
    raster_paths = body.get("raster_paths", [])
    stats = body.get("stats", "mean")

    if not fishnet_path or not Path(fishnet_path).exists():
        raise HTTPException(400, "fishnet_path not found")
    if not raster_paths:
        raise HTTPException(400, "raster_paths required")

    gdf = _gpd.read_file(fishnet_path)
    gdf = run_zonal_stats(gdf, raster_paths, stats=stats)

    out_path = Path(fishnet_path).with_name("fishnet_with_stats.gpkg")
    gdf.to_file(out_path, driver="GPKG")

    # Also export CSV for pipeline consumption
    csv_path = out_path.with_suffix(".csv")
    gdf.drop(columns=["geometry"]).to_csv(csv_path, index=False)

    return {"n_cells": len(gdf), "columns": list(gdf.columns), "csv_path": str(csv_path)}


@app.post("/data/prepare")
async def prepare_data_pipeline(body: dict = Body(...)):
    """Unified pipeline: boundary + rasters → fishnet → zonal stats → CSV.

    Body fields:
        boundary_path: str — path to boundary shapefile/GPKG/GeoJSON
        raster_paths: list[str] — paths to GeoTIFF raster layers
        resolution: float — fishnet cell size (in CRS units, default 100)
        crs: str — target CRS (default EPSG:4326)
        stats: str — zonal stat types, space-separated (default "mean")
        set_as_data: bool — if true, set the result CSV as the project data file (default true)
    """
    if state.project_config is None:
        raise HTTPException(400, "Load a project first.")

    from sparc.data.processing import create_fishnet, clip_to_boundary, run_zonal_stats
    import geopandas as _gpd

    boundary_path = body.get("boundary_path")
    raster_paths = body.get("raster_paths", [])
    resolution = body.get("resolution", 100)
    crs = body.get("crs", "EPSG:4326")
    stats = body.get("stats", "mean")
    set_as_data = body.get("set_as_data", True)

    if not raster_paths:
        raise HTTPException(400, "At least one raster_path is required")

    # Step 1: Determine bounds — from boundary or from rasters
    boundary_gdf = None
    if boundary_path and Path(boundary_path).exists():
        boundary_gdf = _gpd.read_file(boundary_path)
        if boundary_gdf.crs and boundary_gdf.crs.to_string() != crs:
            boundary_gdf = boundary_gdf.to_crs(crs)
        bounds = tuple(boundary_gdf.total_bounds)
    else:
        # Derive bounds from first raster
        try:
            import rasterio
            with rasterio.open(raster_paths[0]) as src:
                bounds = tuple(src.bounds)
                if not crs:
                    crs = str(src.crs)
        except Exception as exc:
            raise HTTPException(400, f"Cannot determine bounds: {exc}")

    # Step 2: Generate fishnet
    gdf = create_fishnet(bounds, resolution, crs)

    # Step 3: Clip to boundary
    if boundary_gdf is not None:
        gdf = clip_to_boundary(gdf, boundary_gdf)

    if len(gdf) == 0:
        raise HTTPException(400, "Fishnet has 0 cells after clipping — check CRS/resolution")

    # Step 4: Run zonal statistics
    valid_rasters = [r for r in raster_paths if Path(r).exists()]
    if not valid_rasters:
        raise HTTPException(400, "None of the raster paths exist")
    gdf = run_zonal_stats(gdf, valid_rasters, stats=stats)

    # Step 5: Save outputs (versioned)
    project_dir = Path(state.project_config["paths"]["project_root"])
    data_dir = project_dir / "data"
    data_dir.mkdir(parents=True, exist_ok=True)

    gpkg_path = data_dir / "fishnet_with_stats.gpkg"
    gdf.to_file(gpkg_path, driver="GPKG")

    # Add coordinate columns before saving
    export_df = gdf.copy()
    export_df["centroid_x"] = gdf.geometry.centroid.x
    export_df["centroid_y"] = gdf.geometry.centroid.y
    export_flat = export_df.drop(columns=["geometry"])

    # Save versioned snapshot
    from sparc.data.versioning import save_versioned
    version_info = save_versioned(
        export_flat,
        data_dir,
        settings={
            "resolution": resolution,
            "crs": crs,
            "stats": stats,
            "n_rasters": len(valid_rasters),
            "boundary": boundary_path,
        },
        description=f"Fishnet processing: {len(valid_rasters)} rasters, "
                    f"resolution={resolution}, CRS={crs}",
    )
    csv_path = version_info["csv_path"]

    # Also save as canonical name for backwards compatibility
    canonical_csv = data_dir / "fishnet_with_stats.csv"
    export_flat.to_csv(canonical_csv, index=False)

    # Step 6: Optionally set as active data
    if set_as_data:
        _set_data_path(str(csv_path))
        _load_data_into_state(state.project_config)

    return {
        "status": "prepared",
        "n_cells": len(gdf),
        "columns": list(gdf.columns),
        "csv_path": str(csv_path),
        "gpkg_path": str(gpkg_path),
        "set_as_data": set_as_data,
        "version": version_info.get("version"),
    }


# ------------------------------------------------------------------
# Report generation
# ------------------------------------------------------------------

@app.post("/report/generate")
async def generate_report(format: str = Query("markdown", regex="^(markdown|json)$")):
    """Generate a pipeline report summarising config, data, and results.

    Returns Markdown by default; JSON structure when ``format=json``.
    """
    import pandas as pd

    sections: list[str] = []
    report_data: dict[str, Any] = {}

    # --- Project info ---
    cfg = state.project_config or {}
    proj = cfg.get("project", {})
    sections.append(f"# SPARC Pipeline Report\n")
    sections.append(f"**Project:** {proj.get('name', 'Untitled')}\n")
    sections.append(f"**Domain:** {proj.get('domain', 'N/A')}\n")
    if proj.get("description"):
        sections.append(f"**Description:** {proj['description']}\n")
    report_data["project"] = proj

    # --- Data overview ---
    if state.data_summary:
        s = state.data_summary
        sections.append("## Data Overview\n")
        sections.append(f"- **Rows:** {s.get('row_count', '?')}")
        sections.append(f"- **Columns:** {s.get('column_count', '?')}")
        if s.get("crs"):
            sections.append(f"- **CRS:** {s['crs']}")
        sections.append("")
        report_data["data_summary"] = s

    # --- Predictors ---
    predictors = cfg.get("predictors", {})
    if isinstance(predictors, list):
        sections.append(f"## Predictors ({len(predictors)})\n")
        sections.append(", ".join(f"`{p}`" for p in predictors))
        sections.append("")
        report_data["predictors"] = predictors
    elif isinstance(predictors, dict):
        base = predictors.get("base_model", [])
        sections.append(f"## Predictors ({len(base)})\n")
        sections.append(", ".join(f"`{p}`" for p in base))
        sections.append("")
        report_data["predictors"] = base

    # --- Causal config ---
    causal = cfg.get("causal", {})
    if causal:
        sections.append("## Causal Configuration\n")
        sections.append(f"- **Estimator:** {causal.get('estimator', 'N/A')}")
        sections.append(f"- **DAG blend weight:** {causal.get('dag_blend_weight', 'N/A')}")
        av = causal.get("actionable_variables", [])
        if av:
            sections.append(f"- **Actionable:** {', '.join(av)}")
        sections.append("")
        report_data["causal"] = causal

    # --- Physics priors ---
    physics = cfg.get("physics", {})
    if physics:
        mc = physics.get("monotone_constraints", {})
        if mc:
            sections.append("## Physics Constraints\n")
            for k, v in mc.items():
                direction = "increasing" if v > 0 else "decreasing" if v < 0 else "none"
                sections.append(f"- `{k}`: {direction}")
            sections.append("")
            report_data["physics"] = physics

    # --- Stage results summary ---
    sections.append("## Results Summary\n")
    sections.append("| Stage | Rows | Metrics |")
    sections.append("|-------|------|---------|")

    for stage_num in [0, 2, 3, 4]:
        result = state.get_result(stage_num)
        if result is None:
            continue
        if isinstance(result, dict) and "summary" in result:
            df = result["summary"]
        elif isinstance(result, pd.DataFrame):
            df = result
        elif isinstance(result, dict) and "rows" in result:
            df = pd.DataFrame(result["rows"])
        else:
            continue

        if hasattr(df, "__len__"):
            n_rows = len(df)
            n_metrics = len(df.columns) if hasattr(df, "columns") else 0
            sections.append(f"| Stage {stage_num} | {n_rows} | {n_metrics} |")
            report_data[f"stage_{stage_num}"] = {"rows": n_rows, "metrics": n_metrics}

    sections.append("")

    # --- Pipeline config ---
    pipeline = cfg.get("pipeline", {})
    if pipeline:
        sections.append("## Pipeline Settings\n")
        sections.append(f"- **Random seed:** {pipeline.get('random_seed', 'N/A')}")
        sections.append(f"- **Spatial folds:** {pipeline.get('n_spatial_folds', 'N/A')}")
        sections.append(f"- **Fast mode:** {pipeline.get('fast_mode', False)}")
        sections.append("")

    md_text = "\n".join(sections)

    if format == "json":
        return {"report": report_data, "markdown": md_text}

    return JSONResponse(content={"markdown": md_text}, media_type="application/json")


@app.post("/report/pdf")
async def generate_pdf_report():
    """Generate a PDF report with embedded plots and return it as a file download."""
    from starlette.responses import Response as StarletteResponse

    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    cfg = state.project_config
    output_dir = cfg.get("paths", {}).get("output_dir")

    # Collect causal results from state
    causal_results = None
    causal_state = state.get_result(3)
    if isinstance(causal_state, dict):
        causal_results = causal_state

    # Collect scenario summary
    scenario_summary = None
    scenario_state = state.get_result(4)
    if isinstance(scenario_state, dict) and "summary" in scenario_state:
        import pandas as pd
        s = scenario_state["summary"]
        if isinstance(s, pd.DataFrame):
            scenario_summary = s.to_dict(orient="records")
        elif isinstance(s, list):
            scenario_summary = s

    try:
        from sparc.report import generate_report_pdf

        pdf_bytes = generate_report_pdf(
            config=cfg,
            data_summary=state.data_summary,
            causal_results=causal_results,
            scenario_summary=scenario_summary,
            output_dir=output_dir,
            registry=state.registry,
        )

        # Also save a copy to the project directory
        project_dir = Path(cfg["paths"]["project_root"])
        dest = project_dir / "sparc_report.pdf"
        if isinstance(pdf_bytes, Path):
            pdf_data = pdf_bytes.read_bytes()
        else:
            pdf_data = pdf_bytes
            dest.write_bytes(pdf_data)

        return StarletteResponse(
            content=pdf_data,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=sparc_report.pdf"},
        )
    except RuntimeError as exc:
        # weasyprint not installed — fall back to HTML
        from sparc.report import generate_report_html

        html_str = generate_report_html(
            config=cfg,
            data_summary=state.data_summary,
            causal_results=causal_results,
            scenario_summary=scenario_summary,
            output_dir=output_dir,
            registry=state.registry,
        )
        # Save HTML to disk
        project_dir = Path(cfg["paths"]["project_root"])
        dest = project_dir / "sparc_report.html"
        dest.write_text(html_str, encoding="utf-8")

        return {
            "status": "html_fallback",
            "message": str(exc),
            "html_path": str(dest),
        }


@app.post("/report/docx")
async def generate_docx_report():
    """Generate a DOCX (Microsoft Word) report and return it as a download.

    DB-only: pulls every figure and table through the run registry / artifacts
    cache via ``sparc.report.generate_report_html`` upstream helpers, then
    composes a Word document from the same structured payload.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise HTTPException(
            503, f"python-docx not installed: {exc}. Run `pip install python-docx`."
        )

    import io as _io

    cfg = state.project_config

    # Reuse the same data-collection logic as /report/pdf.
    causal_results = None
    causal_state = state.get_result(3)
    if isinstance(causal_state, dict):
        causal_results = causal_state

    scenario_summary = None
    scenario_state = state.get_result(4)
    if isinstance(scenario_state, dict) and "summary" in scenario_state:
        import pandas as pd
        s = scenario_state["summary"]
        if isinstance(s, pd.DataFrame):
            scenario_summary = s.to_dict(orient="records")
        elif isinstance(s, list):
            scenario_summary = s

    project_meta = cfg.get("project", {}) if isinstance(cfg, dict) else {}
    project_name = project_meta.get("name", "SPARC Project")

    # Build the document.
    doc = Document()

    title = doc.add_heading(f"SPARC Report — {project_name}", level=0)
    for run in title.runs:
        run.font.size = Pt(18)

    if project_meta.get("description"):
        doc.add_paragraph(str(project_meta["description"]))

    # Data summary
    doc.add_heading("Data Summary", level=1)
    if isinstance(state.data_summary, dict) and state.data_summary:
        for key, value in state.data_summary.items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
    else:
        doc.add_paragraph("No data summary captured for this run.")

    # Causal results
    doc.add_heading("Causal Inference", level=1)
    if isinstance(causal_results, dict) and causal_results:
        for key, value in list(causal_results.items())[:25]:
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
    else:
        doc.add_paragraph("Causal inference results are not yet available.")

    # Scenario summary as a table
    doc.add_heading("Scenario Summary", level=1)
    if scenario_summary:
        cols = list({k for row in scenario_summary for k in row.keys()})
        cols = cols[:8]  # cap to keep table readable
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = str(c)
        for row in scenario_summary[:50]:
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(row.get(c, ""))
    else:
        doc.add_paragraph("Scenario results are not yet available.")

    # Embed registered PNG figures from artifacts.db.
    doc.add_heading("Figures", level=1)
    figures_added = 0
    if state.registry is not None:
        try:
            from sparc.report.figures import (
                FigureRenderError, render_for_artifact,
            )
            for stage_key, stage_obj in list(state.registry.manifest.stages.items()):
                for art_id, art in list(stage_obj.artifacts.items()):
                    if getattr(art, "partial", False):
                        continue
                    try:
                        png = render_for_artifact(
                            str(stage_key), art_id,
                            registry=state.registry, dpi=120,
                        )
                    except FigureRenderError:
                        continue
                    except Exception:
                        continue
                    doc.add_paragraph(f"Stage {stage_key} — {art_id}",
                                      style="Intense Quote")
                    bio = _io.BytesIO(png)
                    try:
                        doc.add_picture(bio, width=Inches(5.5))
                        figures_added += 1
                    except Exception:
                        continue
        except ImportError:
            pass

    if figures_added == 0:
        doc.add_paragraph("No figure renderers produced output for this run.")

    # Serialize to bytes — never touches the project directory.
    buf = _io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    return Response(
        content=data,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": "attachment; filename=sparc_report.docx"},
    )


# ------------------------------------------------------------------
# DAG endpoints
# ------------------------------------------------------------------

@app.get("/dag")
async def get_dag():
    """Return the project's DAG definition."""
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    dag_file = state.project_config.get("causal", {}).get("dag_file")
    if not dag_file or not Path(dag_file).exists():
        return {"nodes": [], "edges": []}

    from sparc.causal.dag_definition import load_dag
    dag = load_dag(dag_file)
    return dag


@app.post("/dag/validate")
async def validate_dag(dag: dict):
    """Validate a DAG definition without saving it."""
    from sparc.causal.dag_definition import dag_to_networkx
    import networkx as nx

    try:
        G = dag_to_networkx(dag)
        is_dag = nx.is_directed_acyclic_graph(G)
        return {
            "valid": is_dag,
            "n_nodes": G.number_of_nodes(),
            "n_edges": G.number_of_edges(),
            "error": None if is_dag else "Graph contains cycles",
        }
    except Exception as exc:
        return {"valid": False, "error": str(exc)}


@app.get("/dag/mc3_result")
async def get_mc3_result():
    """Return pending MC³ edge-inclusion probabilities for DAG approval."""
    if state.pending_mc3 is None:
        raise HTTPException(404, "No MC³ result pending")
    return state.pending_mc3


@app.post("/dag/approve")
async def approve_dag():
    """Approve the discovered DAG and unblock the pipeline."""
    if state.pending_mc3 is None:
        raise HTTPException(400, "No MC³ result pending approval")
    state.dag_approved.set()
    return {"status": "approved"}


@app.post("/dag/reject")
async def reject_dag():
    """Reject the discovered DAG and cancel the pipeline."""
    if state.pending_mc3 is None:
        raise HTTPException(400, "No MC³ result pending approval")
    # Unblock the gate — the pipeline will continue (we clear pending_mc3
    # but the pipeline thread checks is_running which cancel sets to false)
    state.pending_mc3 = None
    state.dag_approved.set()
    return {"status": "rejected"}


# ------------------------------------------------------------------
# Artifact download
# ------------------------------------------------------------------

# ------------------------------------------------------------------
# Run registry endpoints — see definitions earlier in the file
# (placed before /results/{stage} parameterized route to avoid the
# integer-only path converter swallowing /results/manifest paths).
# ------------------------------------------------------------------

@app.get("/results/availability")
async def get_results_availability():
    """Report which result artifacts are present in artifacts.db (DB-only).

    Returns ``{endpoint_path: {available, source}}`` so the desktop can grey
    out tabs that have no underlying data.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    # Build set of (stage, artifact_id) currently registered.
    db_artifacts: set[tuple[str, str]] = set()
    if state.registry is not None:
        try:
            man = state.registry.manifest
            for stage_obj in man.stages.values():
                for art in stage_obj.artifacts.values():
                    if not art.partial:
                        db_artifacts.add((str(stage_obj.stage), art.id))
        except Exception:
            db_artifacts = set()

    # Endpoint -> primary artifact lookup.  None means "any of a family".
    endpoints: dict[str, tuple[str, str] | None] = {
        "/results/correlogram": ("0", "correlogram_results"),
        "/results/gwen": ("1", "gwen_variable_importance"),
        "/results/model_performance": ("2", "ensemble_results"),
        "/results/spatial_cv/predictions": ("2", "spatial_cv_predictions"),
        "/results/neural_pdp": None,  # family: v2_neural_pdp::*
        "/results/pdp_curves": None,  # family
        "/results/causal": ("3", "scenario_coefficients"),
        "/results/causal/diagnostics": ("3", "causal_diagnostics"),
        "/results/causal/dose_response": ("3", "dose_response_curves"),
        "/results/causal/cate_map": ("3", "cate_summary"),
        "/results/scenarios/nuts_summary": ("3", "nuts_summary"),
        "/results/scenarios/detail": ("4", "scenario_results"),
        "/dag/mc3_result": ("3", "mc3_summary"),
    }

    out: dict[str, dict] = {}
    has_neural_pdp = any(s == "2" and a.startswith("v2_neural_pdp::") for s, a in db_artifacts)
    has_gwrf = (("2", "gwrf_condition_curves") in db_artifacts) or \
               (("3", "gwrf_condition_curves") in db_artifacts)

    for endpoint, key in endpoints.items():
        if endpoint == "/results/neural_pdp":
            available = has_neural_pdp
            source = "artifacts.db:2:v2_neural_pdp::*" if available else ""
        elif endpoint == "/results/pdp_curves":
            available = has_neural_pdp or has_gwrf
            source = "artifacts.db (neural_pdp+gwrf)" if available else ""
        else:
            stage_id, art_id = key  # type: ignore[misc]
            available = (stage_id, art_id) in db_artifacts
            source = f"artifacts.db:{stage_id}:{art_id}" if available else ""
        out[endpoint] = {"available": available, "source": source}

    out["_project"] = {
        "available": True,
        "source": state.project_path,
    }
    return out


@app.get("/results/artifacts")
async def list_artifacts():
    """List every registered artifact across all stages (DB-only).

    Iterates ``state.registry.manifest.stages`` rather than walking the
    filesystem. Each artifact is annotated with a ``download_url`` pointing
    at the canonical ``/artifacts/{stage}/{id}`` native renderer.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(404, "No active run registry")

    stage_labels = {
        "0": "Stage 0 — Correlogram",
        "1": "Stage 1 — GWEN",
        "2": "Stage 2 — Spatial CV",
        "3": "Stage 3 — Causal",
        "4": "Stage 4 — Scenarios",
    }

    artifacts: list[dict] = []
    for stage_key, stage_obj in list(state.registry.manifest.stages.items()):
        for art_id, art in list(stage_obj.artifacts.items()):
            if getattr(art, "partial", False):
                continue
            kind = getattr(art, "kind", "")
            ext = {"table": "csv", "struct": "json", "blob": "bin"}.get(kind, "bin")
            artifacts.append({
                "stage": stage_key,
                "stage_label": stage_labels.get(str(stage_key), f"Stage {stage_key}"),
                "artifact_id": art_id,
                "filename": f"{art_id}.{ext}",
                "kind": kind,
                "extension": f".{ext}",
                "size_bytes": getattr(art, "size_bytes", None),
                "producer": getattr(art, "producer", None),
                "download_url": f"/artifacts/{stage_key}/{art_id}",
            })

    return {"artifacts": artifacts, "total": len(artifacts)}


@app.get("/results/download/{stage}/{file_path:path}")
async def download_artifact(stage: str, file_path: str):
    """Download a registered artifact in its native (or requested) format.

    ``file_path`` is parsed as ``<artifact_id>[.<ext>]`` so legacy frontends
    that still build URLs like ``/results/download/2/predictions.csv`` keep
    working. The ext (``csv|json|geojson|png``) routes through the matching
    ``render_*`` helper; an empty ext falls back to the artifact's native
    rendering. All bytes come from ``artifacts.db``.
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(404, "No active run registry")

    from sparc.registry.run_registry import set_active_registry, get_active_registry
    from sparc.report.render import (
        render_native, render_csv, render_json, render_geojson, RenderError,
    )

    # Strip any leading directory components a legacy caller may have included.
    leaf = file_path.rsplit("/", 1)[-1]
    if "." in leaf:
        artifact_id, _, ext = leaf.rpartition(".")
        ext = ext.lower()
    else:
        artifact_id, ext = leaf, ""

    prev = get_active_registry()
    set_active_registry(state.registry)
    try:
        try:
            if ext == "csv":
                data = render_csv(str(stage), artifact_id)
                media = "text/csv"
                fname = f"{artifact_id}.csv"
            elif ext == "json":
                data = render_json(str(stage), artifact_id)
                media = "application/json"
                fname = f"{artifact_id}.json"
            elif ext == "geojson":
                data = render_geojson(str(stage), artifact_id)
                media = "application/geo+json"
                fname = f"{artifact_id}.geojson"
            elif ext == "png":
                from sparc.report.figures import (
                    FigureRenderError, render_for_artifact,
                )
                try:
                    data = render_for_artifact(
                        str(stage), artifact_id, registry=state.registry,
                    )
                except FigureRenderError as exc:
                    raise HTTPException(404, str(exc))
                media = "image/png"
                fname = f"{artifact_id}.png"
            else:
                data, native_ext = render_native(str(stage), artifact_id)
                media = _RENDER_MIME.get(native_ext, "application/octet-stream")
                fname = f"{artifact_id}.{native_ext}"
        except RenderError as exc:
            raise HTTPException(404, str(exc))
    finally:
        set_active_registry(prev)

    return Response(
        content=data,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@app.get("/results/geopackage")
async def download_geopackage():
    """Merge spatial layers from artifacts.db into a single GeoPackage download.

    Layers come from canonical artifacts only:
      * ``predictions``       ← Stage 2 ``predictions``
      * ``causal_effects``    ← Stage 3 ``cate_summary``
      * ``scenario_deltas``   ← Stage 4 ``scenario_results``
    """
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")
    if state.registry is None:
        raise HTTPException(404, "No active run registry")

    import geopandas as gpd
    import tempfile
    from sparc.registry.store import ArtifactStore

    store = ArtifactStore(state.registry)
    layer_specs = [
        ("predictions",     "2", "predictions"),
        ("causal_effects",  "3", "cate_summary"),
        ("scenario_deltas", "4", "scenario_results"),
    ]

    layers: dict[str, gpd.GeoDataFrame] = {}
    for layer_name, stage_id, art_id in layer_specs:
        if not store.has(stage_id, art_id):
            continue
        try:
            gdf = store.read_table(stage_id, art_id)
        except Exception:
            continue
        if not isinstance(gdf, gpd.GeoDataFrame):
            if "geometry" in gdf.columns:
                gdf = gpd.GeoDataFrame(gdf, geometry="geometry", crs="EPSG:4326")
            else:
                continue
        if len(gdf) > 0:
            layers[layer_name] = gdf

    if not layers:
        raise HTTPException(404, "No spatial artifacts available in artifacts.db")

    project_name = state.project_config.get("project", {}).get("name", "sparc_results")
    tmp = tempfile.NamedTemporaryFile(
        suffix=".gpkg", prefix=f"{project_name}_", delete=False
    )
    tmp_path = tmp.name
    tmp.close()

    try:
        for layer_name, gdf in layers.items():
            gdf.to_file(tmp_path, layer=layer_name, driver="GPKG")
    except Exception as exc:
        raise HTTPException(500, f"GeoPackage creation failed: {exc}")

    return FileResponse(
        tmp_path,
        filename=f"{project_name}_results.gpkg",
        media_type="application/geopackage+sqlite3",
    )


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _set_data_path(abs_path: str) -> None:
    """Update the data file path in memory *and* persist to project.yml on disk.

    This ensures that pipeline stages that re-read project.yml from disk
    (e.g. correlogram_analysis via ``load_config()``) see the correct path.
    """
    # 1. In-memory config
    state.project_config["data"]["file_path"] = abs_path
    state.project_config["paths"]["raw_csv_path"] = abs_path

    # 2. Raw YAML dict (for /project/config GET)
    if state.raw_project_yaml is not None:
        state.raw_project_yaml.setdefault("data", {})["file_path"] = abs_path

    # 3. Persist to disk so load_config() in subprocesses picks it up
    if state.project_path:
        import yaml
        yml_path = Path(state.project_path)
        if yml_path.is_dir():
            yml_path = yml_path / "project.yml"
        if yml_path.exists() and state.raw_project_yaml is not None:
            with open(yml_path, "w", encoding="utf-8") as fh:
                yaml.dump(state.raw_project_yaml, fh,
                          default_flow_style=False, sort_keys=False)


def _load_data_into_state(config: dict) -> None:
    """Load and preprocess the project's CSV into state.data."""
    from sparc.data.data_utils import load_and_preprocess_data

    data_path = config["data"]["file_path"]
    if not os.path.exists(data_path):
        return

    try:
        gdf = load_and_preprocess_data(
            raw_data_path=data_path,
            identifier_col=config["variables"]["identifier"],
            target_col=config["variables"]["target"],
            coords_cols=config["variables"]["coordinates"],
            predictor_cols=config["predictors"]["base_model"],
            initial_crs=config["crs"]["initial"],
            target_crs=config["crs"]["target_projected"],
            output_dir=config.get("output", {}).get("base_dir"),
        )
        state.data = gdf
        _compute_summary(state)
    except Exception as exc:
        print(f"Warning: data pre-load failed: {exc}")


def _compute_summary(st: ServerState) -> dict:
    """Build a compact data summary for the LLM and the UI."""
    if st.data is None:
        return {}
    import pandas as pd

    df = st.data
    numeric = df.select_dtypes(include="number")

    summary = {
        "row_count": len(df),
        "column_count": len(df.columns),
        "columns": list(df.columns),
        "dtypes": {c: str(df[c].dtype) for c in df.columns},
        "numeric_summary": {
            c: {
                "mean": float(numeric[c].mean()),
                "median": float(numeric[c].median()),
                "std": float(numeric[c].std()),
                "min": float(numeric[c].min()),
                "max": float(numeric[c].max()),
            }
            for c in numeric.columns
        },
    }

    # CRS info if available
    if hasattr(df, "crs") and df.crs is not None:
        summary["crs"] = str(df.crs)

    # Bounding box
    if hasattr(df, "total_bounds"):
        bounds = df.total_bounds
        summary["bbox"] = {
            "minx": float(bounds[0]),
            "miny": float(bounds[1]),
            "maxx": float(bounds[2]),
            "maxy": float(bounds[3]),
        }

    st.data_summary = summary
    return summary


def _try_load_from_disk(stage: int) -> Any:
    """Attempt to load stage results from the output directory."""
    if state.project_config is None:
        return None

    from sparc.run.pipeline_paths import PipelinePaths
    import pandas as pd

    try:
        paths = PipelinePaths.from_config(state.project_config)
    except Exception:
        return None

    stage_map = {
        0: paths.stage0_dir,
        1: paths.stage1_dir,
        2: paths.stage2_dir,
        3: paths.stage3_dir,
        4: paths.stage4_dir,
    }
    stage_dir = stage_map.get(stage)
    if stage_dir is None or not stage_dir.exists():
        return None

    # For stage 2, look for V2 neural predictions CSV first (already in original scale)
    if stage == 2:
        neural_predictions = stage_dir / "v2_neural" / "predictions.csv"
        if neural_predictions.exists():
            import geopandas as gpd
            from shapely.geometry import Point
            df = pd.read_csv(neural_predictions)
            if "lon" in df.columns and "lat" in df.columns:
                geometry = [Point(xy) for xy in zip(df["lon"], df["lat"])]
                gdf = gpd.GeoDataFrame(df, geometry=geometry, crs="EPSG:4326")
                return gdf

    # Look for common output files (recursive to catch subdirectories)
    for pattern in ["*.gpkg", "*.geojson", "*.parquet", "*.csv"]:
        files = sorted(stage_dir.rglob(pattern))
        if files:
            f = files[0]
            if f.suffix == ".gpkg":
                import geopandas as gpd
                return gpd.read_file(f)
            elif f.suffix == ".parquet":
                import geopandas as gpd
                return gpd.read_parquet(f)
            elif f.suffix == ".geojson":
                import geopandas as gpd
                return gpd.read_file(f)
            else:
                df = pd.read_csv(f)
                # If CSV has lon/lat, convert to GeoDataFrame
                if "lon" in df.columns and "lat" in df.columns:
                    import geopandas as gpd
                    from shapely.geometry import Point
                    geometry = [Point(xy) for xy in zip(df["lon"], df["lat"])]
                    return gpd.GeoDataFrame(df, geometry=geometry, crs="EPSG:4326")
                return df
    return None


def _to_geojson(data: Any) -> dict:
    """Convert a GeoDataFrame or DataFrame to GeoJSON."""
    import geopandas as gpd

    if isinstance(data, gpd.GeoDataFrame):
        return data.__geo_interface__
    if isinstance(data, dict) and "spatial" in data:
        return data["spatial"].__geo_interface__
    raise HTTPException(400, "Data is not spatial; use format=json")


def _to_json(data: Any) -> Any:
    """Convert a DataFrame or dict to JSON-serializable form."""
    import pandas as pd
    import geopandas as gpd

    if isinstance(data, (pd.DataFrame, gpd.GeoDataFrame)):
        # Drop geometry for JSON serialization
        if isinstance(data, gpd.GeoDataFrame):
            data = pd.DataFrame(data.drop(columns="geometry"))
        return {"rows": data.to_dict(orient="records")}
    if isinstance(data, dict):
        return data
    return {"data": str(data)}

# ---------------------------------------------------------------------------
# Phase 5: Artifact export-on-demand endpoints (db-only architecture)
# ---------------------------------------------------------------------------
# These let the desktop UI download any DB-resident artifact as a file
# (CSV / parquet / JSON / GPKG / joblib / etc.) without keeping
# stage folders on disk.

from fastapi.responses import FileResponse as _FileResponse


def _get_artifact_store():
    if state.registry is None:
        raise HTTPException(503, "No active run/registry")
    from sparc.registry.store import ArtifactStore
    return ArtifactStore(state.registry)


@app.get("/artifacts/{stage}/{artifact_id}/download")
def download_artifact(stage: str, artifact_id: str, fmt: Optional[str] = None):
    """Download an artifact from artifacts.db as a file.

    Query parameters:
      fmt � one of csv, parquet, json, gpkg, geojson, joblib, pkl, npy.
            If omitted, defaults to the artifact's natural format.
    """
    store = _get_artifact_store()
    if not store.has(stage, artifact_id):
        raise HTTPException(404, f"Artifact not found: {stage}/{artifact_id}")
    try:
        path = store.export(stage, artifact_id, fmt=fmt)
    except Exception as exc:
        raise HTTPException(500, f"Export failed: {exc}") from exc
    suffix = path.suffix.lstrip(".") or "bin"
    media = {
        "csv": "text/csv",
        "json": "application/json",
        "parquet": "application/octet-stream",
        "gpkg": "application/geopackage+sqlite3",
        "geojson": "application/geo+json",
        "joblib": "application/octet-stream",
        "pkl": "application/octet-stream",
        "npy": "application/octet-stream",
    }.get(suffix, "application/octet-stream")
    return _FileResponse(
        path,
        media_type=media,
        filename=f"{artifact_id}.{suffix}",
    )


@app.post("/artifacts/{stage}/export")
def export_stage_zip(stage: str):
    """Bundle every artifact for a stage into a downloadable .zip."""
    store = _get_artifact_store()
    import tempfile
    tmpdir = Path(tempfile.mkdtemp(prefix=f"sparc_export_{stage}_"))
    try:
        zip_path = store.export_stage(stage, tmpdir, as_zip=True)
        if not isinstance(zip_path, Path):
            raise RuntimeError("export_stage did not return a zip path")
    except Exception as exc:
        raise HTTPException(500, f"Stage export failed: {exc}") from exc
    return _FileResponse(
        zip_path,
        media_type="application/zip",
        filename=f"stage_{stage}_artifacts.zip",
    )


@app.get("/artifacts/{stage}")
def list_stage_artifacts(stage: str):
    """List all artifacts registered for a stage (id, format, storage_kind)."""
    if state.registry is None:
        raise HTTPException(503, "No active run/registry")
    entries = state.registry.list_for_stage(stage)
    return {
        "stage": str(stage),
        "count": len(entries),
        "artifacts": [
            {
                "artifact_id": e.id,
                "format": e.format,
                "storage_kind": getattr(e, "storage_kind", "legacy_path"),
                "partial": bool(e.partial),
                "producer": getattr(e, "producer", None),
            }
            for e in entries
        ],
    }
