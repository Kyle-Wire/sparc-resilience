"""Report generation routes."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from fastapi import APIRouter, Body, HTTPException, Query
from fastapi.responses import JSONResponse, Response

from sparc.server import deps

router = APIRouter(tags=["report"])


@router.post("/report/generate")
async def generate_report(format: str = Query("markdown", pattern="^(markdown|json)$")):
    """Generate a pipeline report summarising config, data, and results."""
    import pandas as pd

    state = deps.state
    sections: list[str] = []
    report_data: dict[str, Any] = {}

    cfg = state.project_config or {}
    proj = cfg.get("project", {})
    sections.append(f"# SPARC Pipeline Report\n")
    sections.append(f"**Project:** {proj.get('name', 'Untitled')}\n")
    sections.append(f"**Domain:** {proj.get('domain', 'N/A')}\n")
    if proj.get("description"):
        sections.append(f"**Description:** {proj['description']}\n")
    report_data["project"] = proj

    if state.data_summary:
        s = state.data_summary
        sections.append("## Data Overview\n")
        sections.append(f"- **Rows:** {s.get('row_count', '?')}")
        sections.append(f"- **Columns:** {s.get('column_count', '?')}")
        if s.get("crs"):
            sections.append(f"- **CRS:** {s['crs']}")
        sections.append("")
        report_data["data_summary"] = s

    predictors = cfg.get("predictors", {})
    if isinstance(predictors, list):
        sections.append(f"## Predictors ({len(predictors)})\n")
        sections.append(", ".join(f"`{p}`" for p in predictors))
        sections.append("")
        report_data["predictors"] = predictors
    elif isinstance(predictors, dict):
        base = predictors.get("base_model", [])
        sections.append(f"## Predictors ({len(base)})\n")
        sections.append(", ".join(f"`{p}`" for p in base))
        sections.append("")
        report_data["predictors"] = base

    causal = cfg.get("causal", {})
    if causal:
        sections.append("## Causal Configuration\n")
        sections.append(f"- **Estimator:** {causal.get('estimator', 'N/A')}")
        sections.append(f"- **DAG blend weight:** {causal.get('dag_blend_weight', 'N/A')}")
        av = causal.get("actionable_variables", [])
        if av:
            sections.append(f"- **Actionable:** {', '.join(av)}")
        sections.append("")
        report_data["causal"] = causal

    physics = cfg.get("physics", {})
    if physics:
        mc = physics.get("monotone_constraints", {})
        if mc:
            sections.append("## Physics Constraints\n")
            for k, v in mc.items():
                direction = "increasing" if v > 0 else "decreasing" if v < 0 else "none"
                sections.append(f"- `{k}`: {direction}")
            sections.append("")
            report_data["physics"] = physics

    sections.append("## Results Summary\n")
    sections.append("| Stage | Rows | Metrics |")
    sections.append("|-------|------|---------|")
    for stage_num in [0, 2, 3, 4]:
        result = state.get_result(stage_num)
        if result is None:
            continue
        if isinstance(result, dict) and "summary" in result:
            df = result["summary"]
        elif isinstance(result, pd.DataFrame):
            df = result
        elif isinstance(result, dict) and "rows" in result:
            df = pd.DataFrame(result["rows"])
        else:
            continue
        if hasattr(df, "__len__"):
            n_rows = len(df)
            n_metrics = len(df.columns) if hasattr(df, "columns") else 0
            sections.append(f"| Stage {stage_num} | {n_rows} | {n_metrics} |")
            report_data[f"stage_{stage_num}"] = {"rows": n_rows, "metrics": n_metrics}
    sections.append("")

    pipeline = cfg.get("pipeline", {})
    if pipeline:
        sections.append("## Pipeline Settings\n")
        sections.append(f"- **Random seed:** {pipeline.get('random_seed', 'N/A')}")
        sections.append(f"- **Spatial folds:** {pipeline.get('n_spatial_folds', 'N/A')}")
        sections.append(f"- **Fast mode:** {pipeline.get('fast_mode', False)}")
        sections.append("")

    md_text = "\n".join(sections)

    if format == "json":
        return {"report": report_data, "markdown": md_text}

    return JSONResponse(content={"markdown": md_text}, media_type="application/json")


@router.post("/report/pdf")
async def generate_pdf_report():
    """Generate a PDF report with embedded plots and return it as a file download."""
    from starlette.responses import Response as StarletteResponse

    state = deps.state
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    cfg = state.project_config
    output_dir = cfg.get("paths", {}).get("output_dir")

    causal_results = None
    causal_state = state.get_result(3)
    if isinstance(causal_state, dict):
        causal_results = causal_state

    scenario_summary = None
    scenario_state = state.get_result(4)
    if isinstance(scenario_state, dict) and "summary" in scenario_state:
        import pandas as pd
        s = scenario_state["summary"]
        if isinstance(s, pd.DataFrame):
            scenario_summary = s.to_dict(orient="records")
        elif isinstance(s, list):
            scenario_summary = s

    try:
        from sparc.report import generate_report_pdf

        pdf_bytes = generate_report_pdf(
            config=cfg,
            data_summary=state.data_summary,
            causal_results=causal_results,
            scenario_summary=scenario_summary,
            output_dir=output_dir,
            registry=state.registry,
        )

        project_dir = Path(cfg["paths"]["project_root"])
        dest = project_dir / "sparc_report.pdf"
        if isinstance(pdf_bytes, Path):
            pdf_data = pdf_bytes.read_bytes()
        else:
            pdf_data = pdf_bytes
            dest.write_bytes(pdf_data)

        return StarletteResponse(
            content=pdf_data,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=sparc_report.pdf"},
        )
    except RuntimeError as exc:
        from sparc.report import generate_report_html

        html_str = generate_report_html(
            config=cfg,
            data_summary=state.data_summary,
            causal_results=causal_results,
            scenario_summary=scenario_summary,
            output_dir=output_dir,
            registry=state.registry,
        )
        project_dir = Path(cfg["paths"]["project_root"])
        dest = project_dir / "sparc_report.html"
        dest.write_text(html_str, encoding="utf-8")
        return {
            "status": "html_fallback",
            "message": str(exc),
            "html_path": str(dest),
        }


@router.post("/report/docx")
async def generate_docx_report():
    """Generate a DOCX (Microsoft Word) report and return it as a download."""
    import io as _io

    state = deps.state
    if state.project_config is None:
        raise HTTPException(400, "No project loaded")

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise HTTPException(
            503, f"python-docx not installed: {exc}. Run `pip install python-docx`."
        )

    cfg = state.project_config

    causal_results = None
    causal_state = state.get_result(3)
    if isinstance(causal_state, dict):
        causal_results = causal_state

    scenario_summary = None
    scenario_state = state.get_result(4)
    if isinstance(scenario_state, dict) and "summary" in scenario_state:
        import pandas as pd
        s = scenario_state["summary"]
        if isinstance(s, pd.DataFrame):
            scenario_summary = s.to_dict(orient="records")
        elif isinstance(s, list):
            scenario_summary = s

    project_meta = cfg.get("project", {}) if isinstance(cfg, dict) else {}
    project_name = project_meta.get("name", "SPARC Project")

    doc = Document()
    title = doc.add_heading(f"SPARC Report — {project_name}", level=0)
    for run in title.runs:
        run.font.size = Pt(18)

    if project_meta.get("description"):
        doc.add_paragraph(str(project_meta["description"]))

    doc.add_heading("Data Summary", level=1)
    if isinstance(state.data_summary, dict) and state.data_summary:
        for key, value in state.data_summary.items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
    else:
        doc.add_paragraph("No data summary captured for this run.")

    doc.add_heading("Causal Inference", level=1)
    if isinstance(causal_results, dict) and causal_results:
        for key, value in list(causal_results.items())[:25]:
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
    else:
        doc.add_paragraph("Causal inference results are not yet available.")

    doc.add_heading("Scenario Summary", level=1)
    if scenario_summary:
        cols = list({k for row in scenario_summary for k in row.keys()})
        cols = cols[:8]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = str(c)
        for row in scenario_summary[:50]:
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(row.get(c, ""))
    else:
        doc.add_paragraph("Scenario results are not yet available.")

    doc.add_heading("Figures", level=1)
    figures_added = 0
    if state.registry is not None:
        try:
            from sparc.report.figures import FigureRenderError, render_for_artifact

            for stage_key, stage_obj in list(state.registry.manifest.stages.items()):
                for art_id, art in list(stage_obj.artifacts.items()):
                    if getattr(art, "partial", False):
                        continue
                    try:
                        png = render_for_artifact(
                            str(stage_key), art_id,
                            registry=state.registry, dpi=120,
                        )
                    except (FigureRenderError, Exception):
                        continue
                    doc.add_paragraph(
                        f"Stage {stage_key} — {art_id}", style="Intense Quote"
                    )
                    bio = _io.BytesIO(png)
                    try:
                        doc.add_picture(bio, width=Inches(5.5))
                        figures_added += 1
                    except Exception:
                        continue
        except ImportError:
            pass

    if figures_added == 0:
        doc.add_paragraph("No figure renderers produced output for this run.")

    buf = _io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    return Response(
        content=data,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": "attachment; filename=sparc_report.docx"},
    )


@router.post("/report/standalone")
async def post_report_standalone(payload: dict | None = Body(default=None)):
    """Bundle the active run into a single self-contained HTML snapshot."""
    from starlette.responses import Response as StarletteResponse

    from sparc.report.standalone_html import build_standalone_html

    state = deps.state
    payload = payload or {}
    run_dir = payload.get("run_dir")
    if not run_dir:
        if state.project_config is None:
            raise HTTPException(400, "No run_dir and no project loaded")
        from sparc.run.pipeline_paths import PipelinePaths

        run_dir = str(PipelinePaths.from_config(state.project_config).output_dir)
    project_name = payload.get("project_name")
    if not project_name and state.project_config is not None:
        project_name = (state.project_config.get("project") or {}).get("name")
    chat_history = (
        payload.get("chat_history")
        if isinstance(payload.get("chat_history"), list)
        else None
    )
    html_text = build_standalone_html(
        run_dir, chat_history=chat_history, project_name=project_name
    )
    fname = f"sparc_snapshot_{Path(run_dir).name or 'run'}.html"
    return StarletteResponse(
        content=html_text,
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/report/audience")
async def post_report_audience(
    audience: str = Query(..., pattern="^(technical|planner|public)$"),
    fmt: str = Query("html", pattern="^(md|html|pdf)$"),
    payload: dict | None = Body(default=None),
):
    """Render a SPARC run for a specific audience.

    Query params:
      audience: technical | planner | public
      fmt:      md | html | pdf
    """
    from starlette.responses import Response as StarletteResponse

    from sparc.report.audience import generate_audience_report

    state = deps.state
    payload = payload or {}
    run_dir = payload.get("run_dir")
    if not run_dir:
        if state.project_config is None:
            raise HTTPException(400, "No project loaded and no run_dir provided")
        run_dir = state.project_config.get("paths", {}).get("output_dir")
    if not run_dir:
        raise HTTPException(400, "Cannot resolve run_dir")
    run_path = Path(run_dir)
    if not run_path.exists():
        raise HTTPException(404, f"Run directory not found: {run_dir}")

    try:
        result = generate_audience_report(
            run_dir=run_path,
            config=state.project_config,
            audience=audience,
            fmt=fmt,
        )
    except RuntimeError as exc:
        raise HTTPException(501, str(exc))
    except ValueError as exc:
        raise HTTPException(400, str(exc))

    media = {"md": "text/markdown", "html": "text/html", "pdf": "application/pdf"}[fmt]
    fname = f"sparc_{audience}.{fmt}"
    content = result if isinstance(result, (bytes, bytearray)) else result.encode("utf-8")
    return StarletteResponse(
        content=content,
        media_type=media,
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )
